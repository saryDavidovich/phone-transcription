"""
לשונית 'יצירת תמלול' עבור מנהל המוסד - העלאת קובץ, תמלול ברקע (thread,
לא חוסם את הדפדפן), והורדת התוצאה כקובץ Word. נפרד מהזרימה הרגילה של
תלמידים (routes/ivr.py + services/transcribe.py) - זהו כלי עבודה אישי
של המוסד עצמו, לא נצרך מיתרת תלמיד.
"""
import os
import uuid
import shutil
import tempfile
import threading
import logging
from flask import Blueprint, render_template, request, jsonify, send_file
from flask_login import current_user
from app import db
from models import InstitutionUpload
from routes.institution import institution_login_required

log = logging.getLogger(__name__)

institution_transcribe_bp = Blueprint('institution_transcribe', __name__)

STATIC_TMP_DIR = os.path.join(os.path.dirname(__file__), '..', 'static', 'fax_tmp')


@institution_transcribe_bp.route('/institution/transcribe')
@institution_login_required
def transcribe_tab():
    uploads = InstitutionUpload.query.filter_by(institution_id=current_user.id).order_by(
        InstitutionUpload.created_at.desc()
    ).limit(30).all()
    return render_template('institution/transcribe.html', uploads=uploads)


@institution_transcribe_bp.route('/institution/transcribe/upload', methods=['POST'])
@institution_login_required
def upload():
    file = request.files.get('audio_file')
    if not file or not file.filename:
        return jsonify({'error': 'יש לבחור קובץ'}), 400

    tier = request.form.get('tier', 'gemini')
    language = request.form.get('language', 'he')
    output_language = request.form.get('output_language', 'he')

    os.makedirs(STATIC_TMP_DIR, exist_ok=True)
    ext = os.path.splitext(file.filename)[1] or '.wav'
    filename = f"inst_{uuid.uuid4().hex}{ext}"
    dest_path = os.path.join(STATIC_TMP_DIR, filename)
    file.save(dest_path)

    record = InstitutionUpload(
        institution_id=current_user.id,
        original_filename=file.filename,
        tier=tier,
        status='processing',
    )
    db.session.add(record)
    db.session.commit()

    base_url = os.environ.get('APP_BASE_URL', os.environ.get('APP_URL', '')).rstrip('/')
    rec_url = f"{base_url}/static/fax_tmp/{filename}"

    from app import app as flask_app
    t = threading.Thread(
        target=_process_upload,
        args=(flask_app, record.id, rec_url, tier, language, output_language),
        daemon=True,
    )
    t.start()

    return jsonify({'uploadId': record.id})


def _process_upload(flask_app, upload_id, rec_url, tier, language, output_language):
    with flask_app.app_context():
        record = InstitutionUpload.query.get(upload_id)
        try:
            from services.transcribe import _gemini_from_url
            transcript, duration, _ = _gemini_from_url(rec_url, language, output_language)
            if not transcript:
                record.status = 'error'
                record.error_message = 'התמלול נכשל - לא התקבל טקסט'
                db.session.commit()
                return

            docx_filename = _make_docx(transcript, record.original_filename)
            record.transcript = transcript
            record.docx_filename = docx_filename
            record.status = 'done'
            db.session.commit()
        except Exception as e:
            log.exception(f'Institution upload {upload_id} failed')
            record.status = 'error'
            record.error_message = str(e)[:500]
            db.session.commit()


def _make_docx(transcript, original_filename):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_heading(f'תמלול: {original_filename or ""}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for line in transcript.split('\n'):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    os.makedirs(STATIC_TMP_DIR, exist_ok=True)
    filename = f"transcript_{uuid.uuid4().hex}.docx"
    doc.save(os.path.join(STATIC_TMP_DIR, filename))
    return filename


@institution_transcribe_bp.route('/institution/transcribe/status/<int:upload_id>')
@institution_login_required
def status(upload_id):
    record = InstitutionUpload.query.filter_by(id=upload_id, institution_id=current_user.id).first()
    if not record:
        return jsonify({'error': 'לא נמצא'}), 404
    return jsonify({
        'status': record.status,
        'error': record.error_message,
        'downloadUrl': f'/institution/transcribe/download/{record.id}' if record.status == 'done' else None,
    })


@institution_transcribe_bp.route('/institution/transcribe/download/<int:upload_id>')
@institution_login_required
def download(upload_id):
    record = InstitutionUpload.query.filter_by(id=upload_id, institution_id=current_user.id).first()
    if not record or not record.docx_filename:
        return 'לא נמצא', 404
    path = os.path.join(STATIC_TMP_DIR, record.docx_filename)
    return send_file(path, as_attachment=True, download_name=f"תמלול - {record.original_filename or 'הקלטה'}.docx")
