from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, send_file, make_response
from flask_login import login_user, logout_user, login_required, current_user
from werkzeug.security import check_password_hash
from app import db, login_manager
from models import Customer, Recording, Transaction, Settings, AdminUser, ManagerMessage, CustomerMessage
from datetime import datetime, timedelta
from sqlalchemy import func
import io
import os
import logging

log = logging.getLogger(__name__)

admin_bp = Blueprint('admin', __name__)

@admin_bp.context_processor
def inject_new_messages_count():
    from models import ManagerMessage, CustomerMessage, GeneralInboxMessage
    try:
        from sqlalchemy import or_, and_
        count = ManagerMessage.query.filter(
            ManagerMessage.status == 'new',
            or_(
                and_(ManagerMessage.rec_url.isnot(None), ManagerMessage.rec_url != ''),
                ManagerMessage.source == 'institution_contact',
            )
        ).count()
        count += CustomerMessage.query.filter_by(direction='in', is_read=False).count()
        count += GeneralInboxMessage.query.filter_by(is_read=False).count()
    except Exception:
        count = 0
    return {'new_messages_count': count}

@login_manager.user_loader
def load_user(user_id):
    # מוסדות מתחברים דרך אותו Flask-Login, אבל עם קידומת 'inst-' ב-get_id
    # (ראה models.Institution.get_id) כדי להבדיל מהתחברות מנהל-על רגילה.
    if isinstance(user_id, str) and user_id.startswith('inst-'):
        from models import Institution
        return Institution.query.get(int(user_id.split('-', 1)[1]))
    return AdminUser.query.get(int(user_id))

def get_setting(key, default=''):
    s = Settings.query.filter_by(key=key).first()
    return s.value if s else default

def set_setting(key, value):
    s = Settings.query.filter_by(key=key).first()
    if s:
        s.value = value
        s.updated_at = datetime.utcnow()
    else:
        s = Settings(key=key, value=value)
        db.session.add(s)
    db.session.commit()

@admin_bp.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        user = AdminUser.query.filter_by(username=username).first()
        if user and check_password_hash(user.password_hash, password):
            login_user(user)
            return redirect(url_for('admin.dashboard'))
        flash('שם משתמש או סיסמה שגויים')
    return render_template('admin/login.html')

@admin_bp.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('admin.login'))

@admin_bp.route('/')
@admin_bp.route('/dashboard')
@login_required
def dashboard():
    from datetime import timedelta
    today = datetime.utcnow().date()
    today_start = datetime.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
    month_start = today.replace(day=1)
    one_month_ago = datetime.utcnow() - timedelta(days=30)
    one_year_ago = datetime.utcnow() - timedelta(days=365)

    # סטטיסטיקות כלליות
    stats = {
        'total_customers': Customer.query.count(),
        'active_customers': Customer.query.filter_by(is_blocked=False).count(),
        'blocked_customers': Customer.query.filter_by(is_blocked=True).count(),
        'total_recordings': Recording.query.count(),
        'today_recordings': Recording.query.filter(
            Recording.created_at >= today_start
        ).count(),
        'month_revenue': db.session.query(func.sum(Transaction.amount)).filter(
            Transaction.type == 'charge',
            Transaction.created_at >= month_start
        ).scalar() or 0,
        'total_revenue': db.session.query(func.sum(Transaction.amount)).filter(
            Transaction.type == 'charge'
        ).scalar() or 0,
        'total_balance': db.session.query(func.sum(Customer.balance)).scalar() or 0,
        # לקוחות לא פעילים
        'inactive_month': Customer.query.filter(
            ~Customer.id.in_(
                db.session.query(Recording.customer_id).filter(
                    Recording.created_at >= one_month_ago
                )
            )
        ).count(),
        'inactive_year': Customer.query.filter(
            ~Customer.id.in_(
                db.session.query(Recording.customer_id).filter(
                    Recording.created_at >= one_year_ago
                )
            )
        ).count(),
    }

    # חיובי ארנקות היום — לפי סוג, ערוץ, ושיטת שליחה
    today_debits = Recording.query.filter(
        Recording.created_at >= today_start,
        Recording.status == 'delivered',
        Recording.cost > 0
    ).all()

    # סטטיסטיקות OCR היום
    from models import OcrResult
    today_ocr = OcrResult.query.filter(
        OcrResult.created_at >= today_start,
        OcrResult.status == 'completed'
    ).all()
    ocr_stats = {
        'count': len(today_ocr),
        'total_cost': sum(r.cost or 0 for r in today_ocr),
        'total_chars': sum(r.char_count or 0 for r in today_ocr),
    }

    billing_stats = {
        'total_cost': sum(r.cost or 0 for r in today_debits),
        'count': len(today_debits),
        # לפי סוג תמלול
        'basic_cost': sum(r.cost or 0 for r in today_debits if (r.transcription_tier or 'gemini') not in ('premium', 'video')),
        'premium_cost': sum(r.cost or 0 for r in today_debits if r.transcription_tier == 'premium'),
        'video_cost': sum(r.cost or 0 for r in today_debits if r.transcription_tier == 'video'),
        'basic_count': sum(1 for r in today_debits if (r.transcription_tier or 'gemini') not in ('premium', 'video')),
        'premium_count': sum(1 for r in today_debits if r.transcription_tier == 'premium'),
        'video_count': sum(1 for r in today_debits if r.transcription_tier == 'video'),
        # לפי ערוץ
        'phone_cost': sum(r.cost or 0 for r in today_debits if not (r.call_id or '').startswith('email-')),
        'email_cost': sum(r.cost or 0 for r in today_debits if (r.call_id or '').startswith('email-')),
        'phone_count': sum(1 for r in today_debits if not (r.call_id or '').startswith('email-')),
        'email_count': sum(1 for r in today_debits if (r.call_id or '').startswith('email-')),
        # לפי שיטת שליחה
        'sent_email_count': sum(1 for r in today_debits if r.delivery_method == 'email'),
        'sent_fax_count': sum(1 for r in today_debits if r.delivery_method == 'fax'),
    }

    recent_recordings = Recording.query.order_by(Recording.created_at.desc()).limit(10).all()
    recent_transactions = Transaction.query.order_by(Transaction.created_at.desc()).limit(10).all()

    pending_stats = {
        'recordings': Recording.query.filter_by(status='pending_payment').count(),
        'ocr': OcrResult.query.filter_by(status='pending_payment').count(),
    }
    pending_stats['total'] = pending_stats['recordings'] + pending_stats['ocr']

    return render_template('admin/dashboard.html',
        stats=stats,
        billing_stats=billing_stats,
        ocr_stats=ocr_stats,
        pending_stats=pending_stats,
        recent_recordings=recent_recordings,
        recent_transactions=recent_transactions
    )

@admin_bp.route('/customers')
@login_required
def customers():
    from sqlalchemy import case, or_

    search = request.args.get('q', '')
    page = request.args.get('page', 1, type=int)
    sort = request.args.get('sort', '')

    # תתי-שאילתות מצטברות (לא N+1) - מספר/סכום טעינות (charge) ומספר הקלטות, לכל לקוח
    load_sq = db.session.query(
        Transaction.customer_id.label('cid'),
        func.count(Transaction.id).label('load_count'),
        func.sum(Transaction.amount).label('load_total')
    ).filter(Transaction.type == 'charge').group_by(Transaction.customer_id).subquery()

    rec_sq = db.session.query(
        Recording.customer_id.label('cid'),
        func.count(Recording.id).label('rec_count')
    ).group_by(Recording.customer_id).subquery()

    load_count_col = func.coalesce(load_sq.c.load_count, 0)
    load_total_col = func.coalesce(load_sq.c.load_total, 0)
    rec_count_col = func.coalesce(rec_sq.c.rec_count, 0)
    has_contact_col = case(
        (or_(
            (Customer.email.isnot(None)) & (Customer.email != ''),
            (Customer.fax.isnot(None)) & (Customer.fax != '')
        ), 1),
        else_=0
    )

    query = Customer.query \
        .outerjoin(load_sq, Customer.id == load_sq.c.cid) \
        .outerjoin(rec_sq, Customer.id == rec_sq.c.cid) \
        .add_columns(load_count_col, load_total_col, rec_count_col)

    if search:
        query = query.filter(
            Customer.phone.contains(search) |
            Customer.name.contains(search) |
            Customer.email.contains(search)
        )

    sort_map = {
        'contact': [has_contact_col.desc(), Customer.created_at.desc()],
        'ever_loaded': [(load_count_col > 0).desc(), load_total_col.desc()],
        'balance': [Customer.balance.desc()],
        'most_used': [rec_count_col.desc()],
        'load_count': [load_count_col.desc()],
    }
    query = query.order_by(*sort_map.get(sort, [Customer.created_at.desc()]))

    paginated = query.paginate(page=page, per_page=50)

    customers_list = [row[0] for row in paginated.items]
    recording_counts = {row[0].id: row[3] for row in paginated.items}
    load_counts = {row[0].id: row[1] for row in paginated.items}
    load_totals = {row[0].id: float(row[2]) for row in paginated.items}

    return render_template('admin/customers.html',
        customers=paginated, customers_list=customers_list, search=search,
        recording_counts=recording_counts, load_counts=load_counts, load_totals=load_totals,
        current_sort=sort)

@admin_bp.route('/customers/add', methods=['GET', 'POST'])
@login_required
def add_customer():
    if request.method == 'POST':
        phone = request.form.get('phone', '').strip()
        name = request.form.get('name', '').strip()
        email = request.form.get('email', '').strip()
        fax = request.form.get('fax', '').strip()
        balance = float(request.form.get('balance', 0) or 0)
        delivery_method = request.form.get('delivery_method', 'email')

        if not phone:
            flash('מספר טלפון הוא שדה חובה')
            return render_template('admin/add_customer.html')

        existing = Customer.query.filter_by(phone=phone).first()
        if existing:
            flash('לקוח עם מספר טלפון זה כבר קיים')
            return render_template('admin/add_customer.html')

        customer = Customer(
            phone=phone,
            name=name,
            email=email,
            fax=fax,
            balance=balance,
            delivery_method=delivery_method
        )
        db.session.add(customer)
        db.session.commit()

        if balance > 0:
            txn = Transaction(
                customer_id=customer.id,
                amount=balance,
                type='credit',
                description='יתרה ראשונית'
            )
            db.session.add(txn)
            db.session.commit()

        flash(f'לקוח {phone} נוסף בהצלחה')
        return redirect(url_for('admin.customer_detail', id=customer.id))

    return render_template('admin/add_customer.html')

@admin_bp.route('/customers/export/excel')
@login_required
def export_customers_excel():
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'לקוחות'

    headers = ['ID', 'טלפון', 'שם', 'מייל', 'פקס', 'יתרה', 'שיטת משלוח', 'חסום', 'תאריך הצטרפות']
    ws.append(headers)

    for cell in ws[1]:
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = PatternFill(start_color='2563EB', end_color='2563EB', fill_type='solid')
        cell.alignment = Alignment(horizontal='center')

    customers = Customer.query.order_by(Customer.created_at.desc()).all()
    for c in customers:
        ws.append([
            c.id,
            c.phone,
            c.name or '',
            c.email or '',
            c.fax or '',
            round(c.balance, 2),
            c.delivery_method or '',
            'כן' if c.is_blocked else 'לא',
            c.created_at.strftime('%d/%m/%Y %H:%M') if c.created_at else ''
        ])

    for col in ws.columns:
        max_length = max(len(str(cell.value or '')) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = max(max_length + 2, 12)

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f'customers_{datetime.now().strftime("%Y%m%d")}.xlsx'
    )

@admin_bp.route('/customers/<int:id>')
@login_required
def customer_detail(id):
    from models import OcrResult, ConversationThread
    customer = Customer.query.get_or_404(id)
    recordings = Recording.query.filter_by(customer_id=id).order_by(Recording.created_at.desc()).all()
    transactions = Transaction.query.filter_by(customer_id=id).order_by(Transaction.created_at.desc()).all()
    ocr_results = OcrResult.query.filter_by(customer_id=id).order_by(OcrResult.created_at.desc()).all()
    threads = ConversationThread.query.filter_by(customer_id=id).order_by(ConversationThread.created_at.desc()).all()
    # צפייה בעמוד מסמנת הודעות נכנסות כ"נקראו" בכל השיחות - כדי שהתראה בעמוד הודעות למנהל תיעלם
    unread = [m for t in threads for m in t.messages if m.direction == 'in' and not m.is_read]
    if unread:
        for m in unread:
            m.is_read = True
        db.session.commit()
    return render_template('admin/customer_detail.html',
        customer=customer, recordings=recordings, transactions=transactions,
        ocr_results=ocr_results, threads=threads, timedelta=timedelta)

@admin_bp.route('/customers/<int:id>/block', methods=['POST'])
@login_required
def block_customer(id):
    customer = Customer.query.get_or_404(id)
    customer.is_blocked = not customer.is_blocked
    db.session.commit()
    status = 'נחסם' if customer.is_blocked else 'בוטל חסם'
    flash(f'לקוח {status} בהצלחה')
    return redirect(url_for('admin.customer_detail', id=id))

@admin_bp.route('/customers/<int:id>/delete', methods=['POST'])
@login_required
def delete_customer(id):
    from models import OcrResult
    customer = Customer.query.get_or_404(id)
    Transaction.query.filter_by(customer_id=id).delete()
    Recording.query.filter_by(customer_id=id).delete()
    OcrResult.query.filter_by(customer_id=id).delete()
    db.session.delete(customer)
    db.session.commit()
    flash(f'לקוח {customer.phone} נמחק בהצלחה')
    return redirect(url_for('admin.customers'))

@admin_bp.route('/customers/<int:id>/credit', methods=['POST'])
@login_required
def credit_customer(id):
    customer = Customer.query.get_or_404(id)
    amount = float(request.form.get('amount', 0))
    reason = request.form.get('reason', 'זיכוי ידני')
    if amount > 0:
        customer.balance = round(customer.balance + amount, 2)
        txn = Transaction(
            customer_id=id,
            amount=amount,
            type='credit',
            description=reason
        )
        db.session.add(txn)
        db.session.commit()
        flash(f'לקוח זוכה ב-{amount:.2f} ₪')
    return redirect(url_for('admin.customer_detail', id=id))

@admin_bp.route('/customers/<int:id>/charge', methods=['POST'])
@login_required
def charge_customer(id):
    customer = Customer.query.get_or_404(id)
    amount = float(request.form.get('amount', 0))
    reason = request.form.get('reason', 'חיוב ידני')
    if amount > 0:
        customer.balance = round(customer.balance - amount, 2)
        txn = Transaction(
            customer_id=id,
            amount=-amount,
            type='debit',
            description=reason
        )
        db.session.add(txn)
        db.session.commit()
        flash(f'לקוח חויב ב-{amount:.2f} ₪')
    return redirect(url_for('admin.customer_detail', id=id))

@admin_bp.route('/customers/<int:id>/update', methods=['POST'])
@login_required
def update_customer(id):
    customer = Customer.query.get_or_404(id)
    customer.name = request.form.get('name', customer.name)
    customer.email = request.form.get('email', customer.email)
    customer.fax = request.form.get('fax', customer.fax)
    customer.delivery_method = request.form.get('delivery_method', customer.delivery_method)
    db.session.commit()
    flash('פרטי לקוח עודכנו')
    return redirect(url_for('admin.customer_detail', id=id))

@admin_bp.route('/customers/<int:id>/send-recordings', methods=['POST'])
@login_required
def send_recordings(id):
    import os
    from services.transcribe import _send_email, _send_fax
    customer = Customer.query.get_or_404(id)
    recording_ids = request.form.getlist('recording_ids')
    send_method = request.form.get('send_method', 'email')
    send_to = request.form.get('send_to', '').strip()

    if not recording_ids:
        flash('לא נבחרו הקלטות')
        return redirect(url_for('admin.customer_detail', id=id))

    if not send_to:
        flash('יש להזין כתובת מייל או מספר פקס')
        return redirect(url_for('admin.customer_detail', id=id))

    sent = 0
    for rec_id in recording_ids:
        rec = Recording.query.get(int(rec_id))
        if not rec or not rec.transcript:
            continue
        try:
            rec_url = f'https://www.call2all.co.il/ym/api/DownloadFile?token={os.environ.get("YEMOT_TOKEN","")}&path=ivr2:/recordings/{rec.call_id}.wav'
            if send_method == 'email':
                _send_email(send_to, rec.transcript, customer, rec_url, rec.duration_seconds, source_filename=rec.source_filename, is_premium=(rec.transcription_tier == 'premium'))
            else:
                _send_fax(send_to, rec.transcript, customer, rec.duration_seconds)
            sent += 1
        except Exception as e:
            flash(f'שגיאה בשליחת הקלטה {rec_id}: {e}')

    flash(f'נשלחו {sent} הקלטות בהצלחה')
    return redirect(url_for('admin.customer_detail', id=id))

@admin_bp.route('/customers/<int:id>/send-message', methods=['POST'])
@login_required
def send_customer_message(id):
    """שליחת הודעה ללקוח במייל, מתוך חשבונו בממשק הניהול. אם נשלח thread_id -
    ממשיכים שיחה קיימת (שרשור אמיתי, In-Reply-To/References). אם לא (או
    thread_id='new') - נפתחת שיחה חדשה ונפרדת, שתוצג בממשק בנפרד מהשיחות
    האחרות. השמירה ב-DB מיידית (מהירה); שליחת המייל בפועל רצה ברקע."""
    import threading, uuid as _uuid
    from models import CustomerMessage, ConversationThread
    customer = Customer.query.get_or_404(id)
    body = (request.form.get('body') or '').strip()
    thread_id = request.form.get('thread_id') or ''

    if not body:
        flash('יש להזין תוכן להודעה')
        return redirect(url_for('admin.customer_detail', id=id))

    if not (customer.email or '').strip():
        flash('ללקוח אין כתובת מייל רשומה')
        return redirect(url_for('admin.customer_detail', id=id))

    if thread_id and thread_id != 'new':
        thread = ConversationThread.query.filter_by(id=int(thread_id), customer_id=id).first()
        if not thread:
            flash('השיחה לא נמצאה')
            return redirect(url_for('admin.customer_detail', id=id))
    else:
        thread = ConversationThread(customer_id=id)
        db.session.add(thread)
        db.session.flush()  # כדי לקבל thread.id לפני commit

    # שרשור אמיתי בתוך אותה שיחה: מוצאים את ה-Message-ID של ההודעה היוצאת
    # האחרונה בשיחה הזו (אם יש), כדי לצרף In-Reply-To/References.
    prev_out = CustomerMessage.query.filter_by(thread_id=thread.id, direction='out') \
        .filter(CustomerMessage.message_id.isnot(None)) \
        .order_by(CustomerMessage.created_at.desc()).first()
    domain = (os.environ.get('SENDGRID_FROM_EMAIL', 'noreply@tamlulphone.co.il')).split('@')[-1]
    new_message_id = f"<conv-{thread.id}-{_uuid.uuid4().hex}@{domain}>"

    msg = CustomerMessage(thread_id=thread.id, customer_id=id, direction='out', body=body,
                           is_read=True, message_id=new_message_id)
    db.session.add(msg)
    db.session.commit()

    in_reply_to = prev_out.message_id if prev_out else None
    threading.Thread(
        target=_send_customer_conversation_email,
        args=(customer.email, customer.phone, body, new_message_id, in_reply_to),
        daemon=True,
    ).start()

    flash('ההודעה נשלחת...')
    return redirect(url_for('admin.customer_detail', id=id))


@admin_bp.route('/customers/<int:id>/threads/<int:thread_id>/delete', methods=['POST'])
@login_required
def delete_conversation_thread(id, thread_id):
    from models import ConversationThread, CustomerMessage
    thread = ConversationThread.query.filter_by(id=thread_id, customer_id=id).first_or_404()
    CustomerMessage.query.filter_by(thread_id=thread.id).delete()
    db.session.delete(thread)
    db.session.commit()
    flash('השיחה נמחקה')
    return redirect(url_for('admin.customer_detail', id=id))


def _send_customer_conversation_email(to_email, phone, body, message_id, in_reply_to=None):
    """שולח מייל שיחה ללקוח (מהמנהל). רץ ברקע (thread נפרד) כדי לא לעכב את
    התגובה למנהל. Reply-To מוגדר לכתובת הקליטה הכללית של המערכת. Message-ID
    עצמאי + In-Reply-To/References - כדי שתוכנות מייל ישרשרו נכון בתוך אותה
    שיחה, במקום להציג כל הודעה כמייל נפרד."""
    from routes.email_inbound import TRANSCRIBE_INBOUND_EMAIL, CONVERSATION_MARKER
    try:
        import sendgrid
        from sendgrid.helpers.mail import Mail, Email, ReplyTo, Header

        html = f'''<div dir="rtl" style="font-family:Arial,sans-serif;max-width:600px;margin:auto">
<h2 style="color:#1e3a8a">הודעה מתמלול פון</h2>
<div style="background:#eff6ff;border-right:4px solid #3b82f6;padding:16px;margin:16px 0;border-radius:8px;line-height:1.8;white-space:pre-wrap">
{body}
</div>
<p style="color:#6b7280;font-size:13px;line-height:1.6">
ניתן להשיב למייל זה<br>
לשירותכם<br>
צוות תמלול פון<br>
03-3131795
</p>
</div>'''

        sg = sendgrid.SendGridAPIClient(api_key=os.environ.get('SENDGRID_API_KEY'))
        message = Mail(
            from_email=Email(os.environ.get('SENDGRID_FROM_EMAIL', os.environ.get('GMAIL_USER', '')), 'תמלול פון'),
            to_emails=to_email,
            subject=f'{CONVERSATION_MARKER} {phone}',
            html_content=html,
        )
        message.reply_to = ReplyTo(TRANSCRIBE_INBOUND_EMAIL, 'תמלול פון')
        message.header = Header('Message-ID', message_id)
        if in_reply_to:
            message.header = Header('In-Reply-To', in_reply_to)
            message.header = Header('References', in_reply_to)
        sg.send(message)
    except Exception as e:
        log.error(f"send_customer_conversation_email failed (to={to_email}): {e}")


@admin_bp.route('/customers/<int:id>/send-instructions', methods=['POST'])
@login_required
def admin_send_instructions(id):
    """כפתור ידני בממשק הניהול - שולח ללקוח את אותו מייל הוראות ששלוחה 5-1
    בטלפון שולחת (הוראות לשליחת קבצים לתמלול במייל)."""
    from routes.email_inbound import _send_instructions_email
    customer = Customer.query.get_or_404(id)
    if not (customer.email or '').strip():
        flash('ללקוח אין כתובת מייל רשומה')
        return redirect(url_for('admin.customer_detail', id=id))
    _send_instructions_email(customer.email, customer.phone, customer.name)
    flash('נשלחו הוראות לשליחת קבצים במייל')
    return redirect(url_for('admin.customer_detail', id=id))


@admin_bp.route('/customers/<int:id>/send-handwriting-instructions', methods=['POST'])
@login_required
def admin_send_handwriting_instructions(id):
    """כפתור ידני בממשק הניהול - שולח ללקוח את אותו מייל הוראות ששלוחה 6-1
    בטלפון שולחת (הוראות לשליחת כתב יד לזיהוי OCR)."""
    from routes.email_inbound import _send_handwriting_instructions_email
    customer = Customer.query.get_or_404(id)
    if not (customer.email or '').strip():
        flash('ללקוח אין כתובת מייל רשומה')
        return redirect(url_for('admin.customer_detail', id=id))
    _send_handwriting_instructions_email(customer.email, customer.phone, customer.name)
    flash('נשלחו הוראות לשליחת כתב יד במייל')
    return redirect(url_for('admin.customer_detail', id=id))

@admin_bp.route('/recordings')
@login_required
def recordings():
    page = request.args.get('page', 1, type=int)
    recordings = Recording.query.order_by(Recording.created_at.desc()).paginate(page=page, per_page=50)
    return render_template('admin/recordings.html', recordings=recordings, timedelta=timedelta)

@admin_bp.route('/recordings/bulk-delete', methods=['POST'])
@login_required
def bulk_delete_recordings():
    recording_ids = request.form.getlist('recording_ids')
    if not recording_ids:
        flash('לא נבחרו הקלטות')
    else:
        count = Recording.query.filter(Recording.id.in_(recording_ids)).delete(synchronize_session=False)
        db.session.commit()
        flash(f'{count} הקלטות נמחקו בהצלחה')
    next_url = request.form.get('next') or url_for('admin.recordings')
    return redirect(next_url)

@admin_bp.route('/recordings/<int:id>')
@login_required
def recording_detail(id):
    recording = Recording.query.get_or_404(id)
    return render_template('admin/recording_detail.html', recording=recording)

@admin_bp.route('/recordings/<int:id>/download-audio')
@login_required
def download_audio(id):
    import requests as req
    recording = Recording.query.get_or_404(id)

    yemot_username = __import__('os').environ.get('YEMOT_USERNAME', '')
    yemot_password = __import__('os').environ.get('YEMOT_PASSWORD', '')

    call_id = recording.call_id or ''
    rec_url = f'https://www.call2all.co.il/ym/api/DownloadFile?username={yemot_username}&password={yemot_password}&path=ivr2:/recordings/{call_id}.wav'

    try:
        r = req.get(rec_url, timeout=60)
        r.raise_for_status()
        return send_file(
            io.BytesIO(r.content),
            mimetype='audio/wav',
            as_attachment=True,
            download_name=f'recording_{id}.wav'
        )
    except Exception as e:
        flash(f'שגיאה בהורדת ההקלטה: {e}')
        return redirect(url_for('admin.recording_detail', id=id))

@admin_bp.route('/recordings/<int:id>/play-audio')
@login_required
def play_audio(id):
    import requests as req
    import os
    recording = Recording.query.get_or_404(id)

    yemot_username = os.environ.get('YEMOT_USERNAME', '')
    yemot_password = os.environ.get('YEMOT_PASSWORD', '')

    call_id = recording.call_id or ''
    rec_url = f'https://www.call2all.co.il/ym/api/DownloadFile?username={yemot_username}&password={yemot_password}&path=ivr2:/recordings/{call_id}.wav'

    try:
        r = req.get(rec_url, timeout=60)
        r.raise_for_status()
        response = make_response(r.content)
        response.headers['Content-Type'] = 'audio/wav'
        response.headers['Content-Disposition'] = 'inline'
        return response
    except Exception as e:
        return jsonify({'error': str(e)}), 400

@admin_bp.route('/recordings/<int:id>/download-word')
@login_required
def download_word(id):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    recording = Recording.query.get_or_404(id)
    customer = Customer.query.get(recording.customer_id)

    doc = Document()

    title = doc.add_heading('תמלול שיחה', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph(f'לקוח: {customer.name or customer.phone if customer else ""}').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(f'תאריך: {recording.created_at.strftime("%d/%m/%Y %H:%M") if recording.created_at else ""}').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(f'משך: {recording.duration_seconds // 60} דקות').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph('─' * 50)

    if recording.summary:
        doc.add_heading('סיכום', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph(recording.summary)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_heading('תמלול מלא', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph(recording.transcript or 'אין תמלול')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'transcript_{id}.docx'
    )

@admin_bp.route('/recordings/cleanup', methods=['POST'])
@login_required
def cleanup_old_recordings():
    cutoff = datetime.utcnow() - timedelta(days=30)
    old = Recording.query.filter(Recording.created_at < cutoff).all()
    count = 0
    for rec in old:
        rec.transcript = None
        rec.summary = None
        count += 1
    db.session.commit()
    flash(f'נמחקו תמלולים של {count} הקלטות ישנות')
    return redirect(url_for('admin.recordings'))

@admin_bp.route('/reports')
@login_required
def reports():
    monthly = db.session.query(
        func.to_char(Transaction.created_at, 'YYYY-MM').label('month'),
        func.sum(Transaction.amount).label('revenue')
    ).filter(Transaction.type == 'charge').group_by('month').order_by('month').all()
    top_customers = db.session.query(
        Customer,
        func.sum(Transaction.amount).label('total_spend')
    ).join(Transaction).filter(
        Transaction.type == 'debit'
    ).group_by(Customer.id).order_by(func.sum(Transaction.amount).desc()).limit(10).all()

    return render_template('admin/reports.html',
        monthly=monthly, top_customers=top_customers)

@admin_bp.route('/change-password', methods=['POST'])
@login_required
def change_password():
    from werkzeug.security import check_password_hash, generate_password_hash

    current_password = request.form.get('current_password', '')
    new_password = request.form.get('new_password', '')
    confirm_password = request.form.get('confirm_password', '')

    if not check_password_hash(current_user.password_hash, current_password):
        flash('הסיסמה הנוכחית שגויה')
        return redirect(url_for('admin.settings'))

    if len(new_password) < 8:
        flash('הסיסמה החדשה חייבת להכיל לפחות 8 תווים')
        return redirect(url_for('admin.settings'))

    if new_password != confirm_password:
        flash('הסיסמאות החדשות אינן תואמות')
        return redirect(url_for('admin.settings'))

    current_user.password_hash = generate_password_hash(new_password)
    db.session.commit()
    flash('הסיסמה שונתה בהצלחה')
    return redirect(url_for('admin.settings'))


@admin_bp.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    if request.method == 'POST':
        set_setting('price_per_20min_basic', request.form.get('price_per_20min_basic', '0.90'))
        set_setting('price_per_20min_premium', request.form.get('price_per_20min_premium', '1.90'))
        set_setting('price_per_20min_video', request.form.get('price_per_20min_video', '1.50'))
        set_setting('price_per_1000_chars_ocr', request.form.get('price_per_1000_chars_ocr', '0.10'))
        set_setting('ocr_engine', request.form.get('ocr_engine', 'gemini'))
        set_setting('yemot_token', request.form.get('yemot_token', ''))
        set_setting('yemot_log_path', request.form.get('yemot_log_path', 'ivr2:/199/LogCreditCardOK.ymgr'))
        set_setting('payment_callback_secret', request.form.get('payment_callback_secret', ''))
        set_setting('nedarim_mosad_number', request.form.get('nedarim_mosad_number', ''))
        set_setting('nedarim_api_password', request.form.get('nedarim_api_password', ''))
        set_setting('nedarim_tamal_type', request.form.get('nedarim_tamal_type', '400'))
        # הגדרות בונוס - עד 3 רמות
        for i in range(1, 4):
            set_setting(f'bonus_threshold_{i}', request.form.get(f'bonus_threshold_{i}', ''))
            set_setting(f'bonus_amount_{i}', request.form.get(f'bonus_amount_{i}', ''))
        set_setting('min_balance', request.form.get('min_balance', '5'))
        set_setting('max_recording_seconds', request.form.get('max_recording_seconds', '1800'))
        set_setting('welcome_new', request.form.get('welcome_new', ''))
        set_setting('welcome_returning', request.form.get('welcome_returning', ''))
        set_setting('system_explanation', request.form.get('system_explanation', ''))
        flash('הגדרות נשמרו בהצלחה')
        return redirect(url_for('admin.settings'))

    current_settings = {
        'price_per_20min_basic': get_setting('price_per_20min_basic', '0.90'),
        'price_per_20min_premium': get_setting('price_per_20min_premium', '1.90'),
        'price_per_20min_video': get_setting('price_per_20min_video', '1.50'),
        'price_per_1000_chars_ocr': get_setting('price_per_1000_chars_ocr', '0.10'),
        'ocr_engine': get_setting('ocr_engine', 'gemini'),
        'yemot_token': get_setting('yemot_token', ''),
        'yemot_log_path': get_setting('yemot_log_path', 'ivr2:/199/LogCreditCardOK.ymgr'),
        'payment_callback_secret': get_setting('payment_callback_secret', ''),
        'nedarim_mosad_number': get_setting('nedarim_mosad_number', ''),
        'nedarim_api_password': get_setting('nedarim_api_password', ''),
        'nedarim_tamal_type': get_setting('nedarim_tamal_type', '400'),
        'bonus_thresholds': [
            {'threshold': get_setting(f'bonus_threshold_{i}', ''), 'amount': get_setting(f'bonus_amount_{i}', '')}
            for i in range(1, 4)
        ],
        'min_balance': get_setting('min_balance', '5'),
        'max_recording_seconds': get_setting('max_recording_seconds', '1800'),
        'welcome_new': get_setting('welcome_new', 'שלום וברוכים הבאים למערכת התמלול.'),
        'welcome_returning': get_setting('welcome_returning', 'שלום וברוכים השבים.'),
        'system_explanation': get_setting('system_explanation', 'מערכת התמלול מאפשרת לך להקליט הודעות שיתומללו ויישלחו אליך למייל או לפקס.'),
    }
    return render_template('admin/settings.html', settings=current_settings)

@admin_bp.route('/api/stats')
@login_required
def api_stats():
    last_30 = []
    for i in range(29, -1, -1):
        day = datetime.utcnow().date() - timedelta(days=i)
        revenue = db.session.query(func.sum(Transaction.amount)).filter(
            Transaction.type == 'charge',
            func.date(Transaction.created_at) == day
        ).scalar() or 0
        recordings = Recording.query.filter(
            func.date(Recording.created_at) == day
        ).count()
        last_30.append({'date': str(day), 'revenue': float(revenue), 'recordings': recordings})
    return jsonify(last_30)
@admin_bp.route('/messages')
@login_required
def manager_messages():
    status_filter = request.args.get('status', '')
    # מסננים החוצה הודעות "משוריינות" שמעולם לא הוקלטו בפועל (הלקוח ניתק לפני
    # שהקליט, או הקיש סולמית בלי להשאיר תוכן) - ל-Yemot עדיין דרוש שריון ID
    # מראש כדי שמספר ההודעה בשלוחה יתאים למספר בממשק, אבל אין סיבה להציג
    # רשומות בלי rec_url בפועל כאילו הן הודעות אמיתיות.
    from sqlalchemy import or_, and_
    query = ManagerMessage.query.filter(or_(
        and_(ManagerMessage.rec_url.isnot(None), ManagerMessage.rec_url != ''),
        ManagerMessage.source == 'institution_contact',
    ))
    if status_filter:
        query = query.filter_by(status=status_filter)
    messages = query.order_by(ManagerMessage.created_at.desc()).all()
    unread_replies = CustomerMessage.query.filter_by(direction='in', is_read=False) \
        .order_by(CustomerMessage.created_at.desc()).all()
    from models import GeneralInboxMessage
    unread_inbox_count = GeneralInboxMessage.query.filter_by(is_read=False).count()
    return render_template('admin/manager_messages.html', messages=messages,
        status_filter=status_filter, timedelta=timedelta, unread_replies=unread_replies,
        unread_inbox_count=unread_inbox_count)


@admin_bp.route('/messages/<int:id>/play')
@login_required
def play_manager_message(id):
    import requests as req, io
    msg = ManagerMessage.query.get_or_404(id)
    if not msg.rec_url:
        return jsonify({'error': 'אין הקלטה'}), 404
    try:
        r = req.get(msg.rec_url, timeout=120, stream=False)
        r.raise_for_status()
        return send_file(
            io.BytesIO(r.content),
            mimetype='application/octet-stream',
            as_attachment=False,
            download_name='message.wav'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 400
 
@admin_bp.route('/messages/<int:id>/status', methods=['POST'])
@login_required
def update_message_status(id):
    msg = ManagerMessage.query.get_or_404(id)
    msg.status     = request.form.get('status', msg.status)
    msg.admin_note = request.form.get('admin_note', msg.admin_note)
    db.session.commit()
    flash('סטטוס עודכן בהצלחה')
    return_filter = request.form.get('return_filter', '')
    return redirect(url_for('admin.manager_messages', status=return_filter) if return_filter
                     else url_for('admin.manager_messages'))


@admin_bp.route('/messages/<int:id>/delete', methods=['POST'])
@login_required
def delete_manager_message(id):
    msg = ManagerMessage.query.get_or_404(id)
    db.session.delete(msg)
    db.session.commit()
    flash('הפניה נמחקה')
    return_filter = request.form.get('return_filter', '')
    return redirect(url_for('admin.manager_messages', status=return_filter) if return_filter
                     else url_for('admin.manager_messages'))
@admin_bp.route('/messages/debug')
@login_required
def debug_messages():
    from models import ManagerMessage
    msgs = ManagerMessage.query.all()
    return jsonify([{
        'id': m.id,
        'call_id': m.call_id,
        'rec_url': m.rec_url
    } for m in msgs])

@admin_bp.route('/test-transcribe', methods=['GET', 'POST'])
@login_required
def test_transcribe():
    if request.method == 'POST':
        import os, uuid, tempfile
        from services.transcribe import _gemini_from_url, finalize_alefbot_recording, _send_email, _send_fax

        file = request.files.get('audio_file')
        tier = request.form.get('tier', 'gemini')
        language = request.form.get('language', 'he')
        output_language = request.form.get('output_language', 'he')
        send_method = request.form.get('send_method', 'email')
        send_to = request.form.get('send_to', '')
        customer_id = request.form.get('customer_id', '')

        if not file:
            flash('יש להעלות קובץ')
            return redirect(url_for('admin.test_transcribe'))

        # שמור קובץ זמנית
        tmp = tempfile.NamedTemporaryFile(suffix='.wav', delete=False)
        file.save(tmp.name)
        tmp.close()

        # בנה URL זמני מה-static
        import shutil
        static_dir = os.path.join(os.path.dirname(__file__), '..', 'static', 'fax_tmp')
        os.makedirs(static_dir, exist_ok=True)
        filename = f"test_{uuid.uuid4().hex}.wav"
        dest_path = os.path.join(static_dir, filename)
        shutil.copy(tmp.name, dest_path)
        os.remove(tmp.name)

        base_url = os.environ.get('APP_BASE_URL', '').rstrip('/')
        rec_url = f"{base_url}/static/fax_tmp/{filename}"

        customer = Customer.query.get(int(customer_id)) if customer_id and customer_id.isdigit() else None
        if customer and not send_to:
            send_to = customer.email if send_method == 'email' else customer.fax

        try:
            transcript_variants = None  # רשימת (כותרת, טקסט) להשוואה בתבנית, לטיירים נסיוניים
            if tier == 'gemini':
                transcript, duration, _ = _gemini_from_url(rec_url, language, output_language)
            elif tier == 'gemini_pro_solo':
                # נסיוני - פעימה אחת בלבד עם Gemini 3.1 Pro + פרומפט הגייה אשכנזית-חסידית
                from services.transcribe import _gemini_pro_solo
                transcript, duration = _gemini_pro_solo(rec_url, language, output_language)
            elif tier == 'gemini_review':
                # נסיוני - תמלול מקצועי מבוסס Gemini עם מעבר תיקון שני (במקום אלף בוט)
                from services.transcribe import _gemini_review_pass
                transcript, duration, transcript_raw_first_pass = _gemini_review_pass(rec_url, language, output_language)
                if transcript_raw_first_pass:
                    transcript_variants = [
                        ('לפני תיקון (תמלול ראשוני)', transcript_raw_first_pass),
                        ('אחרי תיקון (גרסה סופית)', transcript),
                    ]
            elif tier == 'gemini_dual_flash':
                # נסיוני - שני תמלולים עצמאיים + מיזוג ע"י Flash (במקום אלף בוט)
                from services.transcribe import _gemini_dual_transcribe_and_merge
                transcript, duration, transcript_a, transcript_b = _gemini_dual_transcribe_and_merge(
                    rec_url, language, output_language, merge_model='gemini-3.5-flash')
                if transcript_a:
                    variants = [('תמלול א\' (עצמאי)', transcript_a)]
                    if transcript_b:
                        variants.append(('תמלול ב\' (עצמאי)', transcript_b))
                    variants.append(('אחרי מיזוג - Flash (גרסה סופית)', transcript))
                    transcript_variants = variants
            elif tier == 'gemini_dual_pro':
                # נסיוני - שני תמלולים עצמאיים + מיזוג ע"י Pro (במקום אלף בוט)
                from services.transcribe import _gemini_dual_transcribe_and_merge
                transcript, duration, transcript_a, transcript_b = _gemini_dual_transcribe_and_merge(
                    rec_url, language, output_language, merge_model='gemini-3.1-pro-preview')
                if transcript_a:
                    variants = [('תמלול א\' (עצמאי)', transcript_a)]
                    if transcript_b:
                        variants.append(('תמלול ב\' (עצמאי)', transcript_b))
                    variants.append(('אחרי מיזוג - Pro (גרסה סופית)', transcript))
                    transcript_variants = variants
            else:
                # AlefBot — שלח ישירות
                from services.transcribe import _alefbot_submit
                job_id, duration = _alefbot_submit(rec_url, f"test_{uuid.uuid4().hex[:8]}")
                flash(f'אלף בוט קיבל את הקובץ — job_id: {job_id}. התמלול יגיע ב-webhook.')
                return redirect(url_for('admin.test_transcribe'))

            if not transcript:
                flash('התמלול נכשל')
                return redirect(url_for('admin.test_transcribe'))

            if send_to and customer:
                if send_method == 'email':
                    _send_email(send_to, transcript, customer, rec_url, duration)
                    flash(f'תמלול נשלח למייל: {send_to}')
                else:
                    _send_fax(send_to, transcript, customer, duration)
                    flash(f'תמלול נשלח לפקס: {send_to}')
            else:
                flash('התמלול הושלם — לא נשלח (לא הוזן לקוח/כתובת)')

            return render_template('admin/test_transcribe.html',
                transcript=transcript,
                duration=duration,
                transcript_variants=transcript_variants,
                customers=Customer.query.order_by(Customer.name).all()
            )

        except Exception as e:
            flash(f'שגיאה: {e}')
            return redirect(url_for('admin.test_transcribe'))

    customers = Customer.query.order_by(Customer.name).all()
    return render_template('admin/test_transcribe.html', transcript=None, transcript_variants=None, customers=customers)
@admin_bp.route('/messages/bulk-status', methods=['POST'])
@login_required
def bulk_update_status():
    data = request.json
    ids = data.get('ids', [])
    status = data.get('status', '')
    for msg_id in ids:
        msg = ManagerMessage.query.get(int(msg_id))
        if msg:
            msg.status = status
    db.session.commit()
    return jsonify({'ok': True})

@admin_bp.route('/messages/bulk-delete', methods=['POST'])
@login_required
def bulk_delete_messages():
    data = request.json
    ids = data.get('ids', [])
    for msg_id in ids:
        msg = ManagerMessage.query.get(int(msg_id))
        if msg:
            db.session.delete(msg)
    db.session.commit()
    return jsonify({'ok': True})


@admin_bp.route('/customers/bulk-action', methods=['POST'])
@login_required
def bulk_customer_action():
    """ניהול לקוחות מרובים — הוספת כסף, אחוזים, סינון"""
    action = request.form.get('action')
    customer_ids = request.form.getlist('customer_ids')

    if not customer_ids:
        flash('לא נבחרו לקוחות')
        return redirect(url_for('admin.customers'))

    customers = Customer.query.filter(Customer.id.in_(customer_ids)).all()

    if action == 'add_amount':
        amount = float(request.form.get('amount', 0) or 0)
        if amount <= 0:
            flash('סכום לא תקין')
            return redirect(url_for('admin.customers'))
        for c in customers:
            c.balance = round(c.balance + amount, 2)
            txn = Transaction(
                customer_id=c.id,
                amount=amount,
                type='credit',
                description=f'זיכוי ידני מנהל — {amount}₪'
            )
            db.session.add(txn)
        db.session.commit()
        flash(f'נוסף ₪{amount} ל-{len(customers)} לקוחות')

    elif action == 'add_percent':
        percent = float(request.form.get('percent', 0) or 0)
        if percent <= 0:
            flash('אחוז לא תקין')
            return redirect(url_for('admin.customers'))
        for c in customers:
            bonus = round(c.balance * percent / 100, 2)
            c.balance = round(c.balance + bonus, 2)
            txn = Transaction(
                customer_id=c.id,
                amount=bonus,
                type='credit',
                description=f'זיכוי {percent}% מנהל — {bonus}₪'
            )
            db.session.add(txn)
        db.session.commit()
        flash(f'נוסף {percent}% ל-{len(customers)} לקוחות')

    elif action == 'delete':
        from models import OcrResult
        count = len(customers)
        for c in customers:
            Transaction.query.filter_by(customer_id=c.id).delete()
            Recording.query.filter_by(customer_id=c.id).delete()
            OcrResult.query.filter_by(customer_id=c.id).delete()
            db.session.delete(c)
        db.session.commit()
        flash(f'{count} לקוחות נמחקו בהצלחה')

    return redirect(url_for('admin.customers'))


@admin_bp.route('/customers-filter')
@login_required
def customers_filter():
    """סינון לקוחות לפי פעילות"""
    from datetime import timedelta
    filter_type = request.args.get('filter', '')
    search = request.args.get('q', '')
    page = request.args.get('page', 1, type=int)

    query = Customer.query

    if filter_type == 'inactive_month':
        one_month_ago = datetime.utcnow() - timedelta(days=30)
        active_ids = db.session.query(Recording.customer_id).filter(
            Recording.created_at >= one_month_ago
        ).subquery()
        query = query.filter(~Customer.id.in_(active_ids))
    elif filter_type == 'inactive_year':
        one_year_ago = datetime.utcnow() - timedelta(days=365)
        active_ids = db.session.query(Recording.customer_id).filter(
            Recording.created_at >= one_year_ago
        ).subquery()
        query = query.filter(~Customer.id.in_(active_ids))

    if search:
        query = query.filter(
            Customer.phone.contains(search) |
            Customer.name.contains(search) |
            Customer.email.contains(search)
        )

    customers = query.order_by(Customer.created_at.desc()).paginate(page=page, per_page=50)

    customer_ids = [c.id for c in customers.items]
    counts = dict(
        db.session.query(Recording.customer_id, func.count(Recording.id))
        .filter(Recording.customer_id.in_(customer_ids))
        .group_by(Recording.customer_id)
        .all()
    ) if customer_ids else {}
    recording_counts = {cid: counts.get(cid, 0) for cid in customer_ids}

    return render_template('admin/customers.html',
        customers=customers,
        search=search,
        active_filter=filter_type,
        recording_counts=recording_counts
    )


@admin_bp.route('/ocr')
@login_required
def ocr_list():
    from models import OcrResult
    page = request.args.get('page', 1, type=int)
    search = request.args.get('q', '')

    query = OcrResult.query.order_by(OcrResult.created_at.desc())
    if search:
        query = query.join(Customer).filter(
            Customer.phone.contains(search) | OcrResult.original_filename.contains(search)
        )

    ocr_results = query.paginate(page=page, per_page=20)
    return render_template('admin/ocr_list.html', ocr_results=ocr_results, search=search)


@admin_bp.route('/ocr/bulk-delete', methods=['POST'])
@login_required
def bulk_delete_ocr():
    from models import OcrResult
    ocr_ids = request.form.getlist('ocr_ids')
    if not ocr_ids:
        flash('לא נבחרו תוצאות OCR')
    else:
        items = OcrResult.query.filter(OcrResult.id.in_(ocr_ids)).all()
        count = len(items)
        for item in items:
            try:
                if item.original_file_path and os.path.exists(item.original_file_path):
                    os.remove(item.original_file_path)
            except Exception:
                pass
            db.session.delete(item)
        db.session.commit()
        flash(f'{count} תוצאות OCR נמחקו בהצלחה')
    next_url = request.form.get('next') or url_for('admin.ocr_list')
    return redirect(next_url)


@admin_bp.route('/ocr/<int:ocr_id>')
@login_required
def ocr_detail(ocr_id):
    from models import OcrResult
    ocr = OcrResult.query.get_or_404(ocr_id)
    return render_template('admin/ocr_detail.html', ocr=ocr)


@admin_bp.route('/ocr/<int:ocr_id>/file')
@login_required
def ocr_file(ocr_id):
    from models import OcrResult
    import os
    ocr = OcrResult.query.get_or_404(ocr_id)
    if not ocr.original_file_path or not os.path.exists(ocr.original_file_path):
        return "הקובץ אינו זמין", 404
    return send_file(ocr.original_file_path, as_attachment=False,
                     download_name=ocr.original_filename)


def _il_day_range_to_utc(start_str, end_str):
    """ממיר טווח תאריכים (YYYY-MM-DD, לפי שעון ישראל) לטווח UTC [start, end)
    להשוואה מול started_at שנשמר ב-UTC. מטפל אוטומטית בשעון קיץ/חורף."""
    from zoneinfo import ZoneInfo
    from datetime import timezone as _timezone
    il_tz = ZoneInfo('Asia/Jerusalem')
    start_local = datetime.strptime(start_str, '%Y-%m-%d').replace(tzinfo=il_tz)
    end_local = (datetime.strptime(end_str, '%Y-%m-%d') + timedelta(days=1)).replace(tzinfo=il_tz)
    start_utc = start_local.astimezone(_timezone.utc).replace(tzinfo=None)
    end_utc = end_local.astimezone(_timezone.utc).replace(tzinfo=None)
    return start_utc, end_utc


@admin_bp.route('/calls-report')
@login_required
def calls_report():
    from zoneinfo import ZoneInfo
    today_il = datetime.now(ZoneInfo('Asia/Jerusalem')).strftime('%Y-%m-%d')
    default_start = request.args.get('start', today_il)
    default_end = request.args.get('end', today_il)
    return render_template('admin/calls_report.html',
        default_start=default_start, default_end=default_end)


@admin_bp.route('/calls-report/export')
@login_required
def calls_report_export():
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from models import CallLog
    from collections import defaultdict

    start_str = request.args.get('start', '')
    end_str = request.args.get('end', '')
    if not start_str or not end_str:
        flash('יש לבחור טווח תאריכים', 'error')
        return redirect(url_for('admin.calls_report'))

    start_utc, end_utc = _il_day_range_to_utc(start_str, end_str)

    calls = CallLog.query.filter(
        CallLog.started_at >= start_utc,
        CallLog.started_at < end_utc
    ).all()

    agg = defaultdict(lambda: {'count': 0, 'total_seconds': 0, 'incomplete': 0})
    for c in calls:
        a = agg[c.phone]
        a['count'] += 1
        if c.duration_seconds is not None:
            a['total_seconds'] += c.duration_seconds
        else:
            a['incomplete'] += 1  # שיחה שעדיין לא נרשם לה סיום (ננטשה/תקלה)

    def fmt_dur(seconds):
        seconds = int(seconds)
        h, rem = divmod(seconds, 3600)
        m, s = divmod(rem, 60)
        return f"{h:02d}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'פילוח שיחות'
    ws.sheet_view.rightToLeft = True

    headers = ['מספר טלפון', 'מספר שיחות', 'זמן שיחה כולל', 'זמן שיחה כולל (שניות)', 'שיחות ללא סיום רשום']
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = PatternFill(start_color='2563EB', end_color='2563EB', fill_type='solid')
        cell.alignment = Alignment(horizontal='center')

    rows = sorted(agg.items(), key=lambda kv: -kv[1]['total_seconds'])
    for phone, a in rows:
        ws.append([phone, a['count'], fmt_dur(a['total_seconds']), a['total_seconds'], a['incomplete']])

    total_calls = sum(a['count'] for a in agg.values())
    total_seconds = sum(a['total_seconds'] for a in agg.values())
    ws.append([])
    ws.append(['סה"כ', total_calls, fmt_dur(total_seconds), total_seconds, ''])
    for cell in ws[ws.max_row]:
        cell.font = Font(bold=True)

    for col in ws.columns:
        max_length = max(len(str(cell.value or '')) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = max(max_length + 2, 12)

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f'calls_report_{start_str}_to_{end_str}.xlsx'
    )


# ==================== הודעה תפוצתית לכלל הלקוחות ====================
# הודעה חד-פעמית שנשלחת לכל הלקוחות עם כתובת מייל רשומה (לא קשור לשיחת
# CustomerMessage/thread ספציפית - זו הודעה כללית, לא שיחה עם לקוח יחיד).
# יש "שליחת ניסיון" לכתובת אחת לפני שמאשרים שליחה אמיתית לכולם - הטופס
# תמיד נטען מחדש עם אותו נושא/תוכן שהוזנו (render_template ולא redirect),
# כדי שכתיבת ניסיון לא תמחק את ההודעה שעדיין לא נשלחה לכולם.

def _broadcast_recipient_emails():
    """כתובות מייל ייחודיות של כל הלקוחות עם מייל רשום, בלי קשר לחסימה/יתרה -
    זו הודעה כללית, לא תמלול בתשלום."""
    rows = Customer.query.filter(Customer.email.isnot(None), Customer.email != '').all()
    seen = set()
    emails = []
    for c in rows:
        e = (c.email or '').strip()
        if e and e.lower() not in seen:
            seen.add(e.lower())
            emails.append(e)
    return emails


def _broadcast_email_html(subject, body):
    """HTML פשוט לפי אותו מיתוג שכבר קיים בהודעות ללקוח (_send_customer_conversation_email) -
    כותרת, תוכן (עם ירידות שורה), וחתימה קבועה."""
    from markupsafe import escape
    safe_body = str(escape(body)).replace('\n', '<br>')
    return f'''<div dir="rtl" style="font-family:Arial,sans-serif;max-width:600px;margin:auto">
<h2 style="color:#1e3a8a">{escape(subject)}</h2>
<div style="background:#eff6ff;border-right:4px solid #3b82f6;padding:16px;margin:16px 0;border-radius:8px;line-height:1.8">
{safe_body}
</div>
<p style="color:#6b7280;font-size:13px;line-height:1.6">
לשירותכם<br>
צוות תמלול פון<br>
03-3131795
</p>
</div>'''


def _send_broadcast_email(to_email, subject, html):
    import sendgrid
    from sendgrid.helpers.mail import Mail, Email
    sg = sendgrid.SendGridAPIClient(api_key=os.environ.get('SENDGRID_API_KEY'))
    message = Mail(
        from_email=Email(os.environ.get('SENDGRID_FROM_EMAIL', os.environ.get('GMAIL_USER', '')), 'תמלול פון'),
        to_emails=to_email,
        subject=subject,
        html_content=html,
    )
    sg.send(message)


def _send_broadcast_bg(subject, html, recipient_emails):
    """רץ ב-thread נפרד: שולח לכל הנמענים ברצף (לא thread לכל מייל - כדי לא
    להציף בהרבה threads כשיש הרבה לקוחות), עם לוג שגיאה לכל כתובת שנכשלה
    בלי לעצור את שאר השליחה."""
    sent, failed = 0, 0
    for email in recipient_emails:
        try:
            _send_broadcast_email(email, subject, html)
            sent += 1
        except Exception as e:
            failed += 1
            log.error(f"broadcast send failed (to={email}): {e}")
    log.info(f"הודעה תפוצתית הסתיימה: נשלחו {sent}, נכשלו {failed} (מתוך {len(recipient_emails)})")


@admin_bp.route('/broadcast', methods=['GET'])
@login_required
def broadcast_message():
    return render_template(
        'admin/broadcast.html',
        subject='', body='', test_email='',
        recipient_count=len(_broadcast_recipient_emails()),
        result=None,
    )


@admin_bp.route('/broadcast/test', methods=['POST'])
@login_required
def broadcast_test():
    subject = (request.form.get('subject') or '').strip()
    body = (request.form.get('body') or '').strip()
    test_email = (request.form.get('test_email') or '').strip()
    recipient_count = len(_broadcast_recipient_emails())

    result = None
    if not subject or not body:
        result = {'ok': False, 'text': 'יש למלא נושא ותוכן לפני שליחת ניסיון.'}
    elif not test_email:
        result = {'ok': False, 'text': 'יש להזין כתובת מייל לבדיקה.'}
    else:
        try:
            html = _broadcast_email_html(subject, body)
            _send_broadcast_email(test_email, f'[בדיקה] {subject}', html)
            result = {'ok': True, 'text': f'נשלח מייל בדיקה אל {test_email}. ההודעה נשארה כאן למטה - אפשר להמשיך לערוך או ללחוץ "שלח לכולם".'}
        except Exception as e:
            log.error(f"broadcast test send failed: {e}")
            result = {'ok': False, 'text': f'שגיאה בשליחת מייל הבדיקה: {e}'}

    return render_template(
        'admin/broadcast.html',
        subject=subject, body=body, test_email=test_email,
        recipient_count=recipient_count, result=result,
    )


@admin_bp.route('/broadcast/send', methods=['POST'])
@login_required
def broadcast_send():
    import threading
    subject = (request.form.get('subject') or '').strip()
    body = (request.form.get('body') or '').strip()
    recipient_emails = _broadcast_recipient_emails()

    result = None
    if not subject or not body:
        result = {'ok': False, 'text': 'יש למלא נושא ותוכן.'}
    elif not recipient_emails:
        result = {'ok': False, 'text': 'אין כרגע אף לקוח עם מייל פעיל לשלוח אליו.'}
    else:
        html = _broadcast_email_html(subject, body)
        threading.Thread(
            target=_send_broadcast_bg,
            args=(subject, html, recipient_emails),
            daemon=True,
        ).start()
        result = {'ok': True, 'text': f'ההודעה נשלחת ברקע ל-{len(recipient_emails)} לקוחות. זה עשוי לקחת כמה דקות (ראה יומן השרת לפירוט תוצאה סופית).'}

    return render_template(
        'admin/broadcast.html',
        subject=subject, body=body, test_email='',
        recipient_count=len(recipient_emails), result=result,
    )


# ==================== תיבה כללית - מיילים ללא זיהוי לקוח ====================
# שרשורים עם גורמים שלא זוהו כלקוח רשום (לא נמצא מספר טלפון בנושא שתואם
# ללקוח קיים) - ראה routes/email_inbound.py: _find_customer_by_phone_anywhere.
# המנהל יכול להשיב במייל (בדיוק כמו שיחה עם לקוח - הצד השני יכול להשיב
# בחזרה וזה ימשיך את אותו שרשור), למחוק שרשור, או לשייך אותו ידנית ללקוח -
# שיוך מעתיק את כל ההודעות לתוך conversation_threads של אותו לקוח ומוחק
# את השרשור מהתיבה הכללית.

def _send_inbox_reply_email(to_email, subject, body, message_id, in_reply_to=None):
    """כמו _send_customer_conversation_email, אבל לשרשור אנונימי (לפי מייל,
    לא לקוח רשום) - נשלח ברקע (thread נפרד) כדי לא לעכב את התגובה למנהל."""
    from routes.email_inbound import TRANSCRIBE_INBOUND_EMAIL
    try:
        import sendgrid
        from sendgrid.helpers.mail import Mail, Email, ReplyTo, Header

        html = f'''<div dir="rtl" style="font-family:Arial,sans-serif;max-width:600px;margin:auto">
<h2 style="color:#1e3a8a">הודעה מתמלול פון</h2>
<div style="background:#eff6ff;border-right:4px solid #3b82f6;padding:16px;margin:16px 0;border-radius:8px;line-height:1.8;white-space:pre-wrap">
{body}
</div>
<p style="color:#6b7280;font-size:13px;line-height:1.6">
ניתן להשיב למייל זה<br>
לשירותכם<br>
צוות תמלול פון<br>
03-3131795
</p>
</div>'''

        sg = sendgrid.SendGridAPIClient(api_key=os.environ.get('SENDGRID_API_KEY'))
        message = Mail(
            from_email=Email(os.environ.get('SENDGRID_FROM_EMAIL', os.environ.get('GMAIL_USER', '')), 'תמלול פון'),
            to_emails=to_email,
            subject=subject,
            html_content=html,
        )
        message.reply_to = ReplyTo(TRANSCRIBE_INBOUND_EMAIL, 'תמלול פון')
        message.header = Header('Message-ID', message_id)
        if in_reply_to:
            message.header = Header('In-Reply-To', in_reply_to)
            message.header = Header('References', in_reply_to)
        sg.send(message)
    except Exception as e:
        log.error(f"send_inbox_reply_email failed (to={to_email}): {e}")


@admin_bp.route('/inbox')
@login_required
def general_inbox():
    from models import GeneralInboxMessage
    threads = GeneralInboxMessage.query.order_by(GeneralInboxMessage.updated_at.desc()).all()
    # שימו לב: בכוונה לא מסמנים כאן כלום כ"נקרא" - רק צפייה בפועל בשרשור ספציפי
    # (inbox_thread_detail למטה) מסמנת אותו. אחרת עצם טעינת הרשימה הזו הייתה
    # מנקה את כל ההתראות במחי יד, גם עבור שרשורים שלא נפתחו בפועל בכלל.
    return render_template('admin/inbox.html', threads=threads)


@admin_bp.route('/inbox/<int:id>')
@login_required
def inbox_thread_detail(id):
    from models import GeneralInboxMessage
    thread = GeneralInboxMessage.query.get_or_404(id)
    if not thread.is_read:
        thread.is_read = True
        db.session.commit()
    return render_template('admin/inbox_thread.html', thread=thread)


@admin_bp.route('/inbox/<int:id>/reply', methods=['POST'])
@login_required
def reply_inbox_thread(id):
    import threading, uuid as _uuid
    from models import GeneralInboxMessage, InboxMessage
    thread = GeneralInboxMessage.query.get_or_404(id)
    body = (request.form.get('body') or '').strip()
    if not body:
        flash('יש להזין תוכן להודעה')
        return redirect(url_for('admin.inbox_thread_detail', id=id))

    prev_out = InboxMessage.query.filter_by(thread_id=thread.id, direction='out') \
        .filter(InboxMessage.message_id.isnot(None)) \
        .order_by(InboxMessage.created_at.desc()).first()
    domain = (os.environ.get('SENDGRID_FROM_EMAIL', 'noreply@tamlulphone.co.il')).split('@')[-1]
    new_message_id = f"<inbox-{thread.id}-{_uuid.uuid4().hex}@{domain}>"

    msg = InboxMessage(thread_id=thread.id, direction='out', body=body, message_id=new_message_id)
    db.session.add(msg)
    thread.is_read = True
    db.session.commit()

    reply_subject = thread.subject or 'תמלול פון'
    if not reply_subject.lower().startswith('re:'):
        reply_subject = f'Re: {reply_subject}'
    in_reply_to = prev_out.message_id if prev_out else None
    threading.Thread(
        target=_send_inbox_reply_email,
        args=(thread.from_email, reply_subject, body, new_message_id, in_reply_to),
        daemon=True,
    ).start()

    flash('ההודעה נשלחת...')
    return redirect(url_for('admin.inbox_thread_detail', id=id))


@admin_bp.route('/inbox/<int:id>/delete', methods=['POST'])
@login_required
def delete_inbox_message(id):
    from models import GeneralInboxMessage, InboxMessage
    thread = GeneralInboxMessage.query.get_or_404(id)
    InboxMessage.query.filter_by(thread_id=thread.id).delete()
    db.session.delete(thread)
    db.session.commit()
    flash('השרשור נמחק')
    return redirect(url_for('admin.general_inbox'))


@admin_bp.route('/inbox/<int:id>/assign', methods=['POST'])
@login_required
def assign_inbox_message(id):
    """משייך שרשור מהתיבה הכללית ללקוח לפי מספר טלפון - מעתיק את כל ההודעות
    (בשני הכיוונים, לפי הסדר) לתוך שיחה חדשה בחשבון שלו, ומוחק את השרשור
    מהתיבה הכללית."""
    from models import GeneralInboxMessage, InboxMessage, CustomerMessage, ConversationThread
    from routes.email_inbound import _normalize_israeli_phone
    thread = GeneralInboxMessage.query.get_or_404(id)
    phone = _normalize_israeli_phone(request.form.get('phone', ''))

    customer = Customer.query.filter_by(phone=phone).first() if phone else None
    if not customer:
        flash('לא נמצא לקוח עם מספר הטלפון הזה')
        return redirect(url_for('admin.inbox_thread_detail', id=id))

    new_thread = ConversationThread(customer_id=customer.id)
    db.session.add(new_thread)
    db.session.flush()

    inbox_messages = InboxMessage.query.filter_by(thread_id=thread.id) \
        .order_by(InboxMessage.created_at.asc()).all()
    for i, m in enumerate(inbox_messages):
        body = m.body
        if i == 0 and thread.subject:
            body = f'[נושא מקורי: {thread.subject}]\n{body}'
        db.session.add(CustomerMessage(
            thread_id=new_thread.id, customer_id=customer.id,
            direction=m.direction, body=body, is_read=True,
        ))

    InboxMessage.query.filter_by(thread_id=thread.id).delete()
    db.session.delete(thread)
    db.session.commit()
    flash(f'השרשור שויך ללקוח {customer.phone}' + (f' ({customer.name})' if customer.name else ''))
    return redirect(url_for('admin.customer_detail', id=customer.id))


# ===================== מוסדות (ניהול מוסד - super-admin) =====================

@admin_bp.route('/institutions')
@login_required
def institutions():
    from models import Institution
    insts = Institution.query.order_by(Institution.created_at.desc()).all()
    return render_template('admin/institutions.html', institutions=insts)


@admin_bp.route('/institutions/create', methods=['POST'])
@login_required
def create_institution():
    import random, string
    from models import Institution
    name = (request.form.get('name') or '').strip()
    email = (request.form.get('email') or '').strip() or None
    phone = (request.form.get('phone') or '').strip() or None
    if not name or not (email or phone):
        flash('יש להזין שם, ולפחות מייל או טלפון')
        return redirect(url_for('admin.institutions'))

    inst = Institution(
        name=name, email=email, phone=phone,
        login_code=''.join(random.choices(string.digits, k=6)),
    )
    db.session.add(inst)
    db.session.commit()
    flash(f'המוסד "{name}" נוצר - קוד כניסה ראשוני: {inst.login_code}')
    return redirect(url_for('admin.institutions'))


@admin_bp.route('/institutions/<int:inst_id>')
@login_required
def institution_detail(inst_id):
    from models import Institution, Customer, InstitutionChargeLog
    inst = Institution.query.get_or_404(inst_id)
    students = Customer.query.filter_by(institution_id=inst.id).all()
    charges = InstitutionChargeLog.query.filter_by(institution_id=inst.id).order_by(InstitutionChargeLog.created_at.desc()).limit(20).all()
    return render_template('admin/institution_detail.html', inst=inst, students=students, charges=charges)


@admin_bp.route('/institutions/<int:inst_id>/adjust-balance', methods=['POST'])
@login_required
def adjust_institution_balance(inst_id):
    from models import Institution
    inst = Institution.query.get_or_404(inst_id)
    try:
        amount = float(request.form.get('amount'))
    except (TypeError, ValueError):
        flash('סכום לא תקין')
        return redirect(url_for('admin.institution_detail', inst_id=inst.id))
    inst.balance = (inst.balance or 0) + amount
    db.session.commit()
    flash('היתרה עודכנה')
    return redirect(url_for('admin.institution_detail', inst_id=inst.id))


@admin_bp.route('/institutions/<int:inst_id>/toggle-block', methods=['POST'])
@login_required
def toggle_institution_block(inst_id):
    from models import Institution
    inst = Institution.query.get_or_404(inst_id)
    inst.is_blocked = not inst.is_blocked
    db.session.commit()
    flash('המוסד נחסם' if inst.is_blocked else 'החסימה בוטלה')
    return redirect(url_for('admin.institution_detail', inst_id=inst.id))


@admin_bp.route('/institutions/<int:inst_id>/regenerate-code', methods=['POST'])
@login_required
def regenerate_institution_code(inst_id):
    import random, string
    from models import Institution
    inst = Institution.query.get_or_404(inst_id)
    inst.login_code = ''.join(random.choices(string.digits, k=6))
    inst.password_hash = None  # חוזר למצב "כניסה ראשונה" עם הקוד החדש
    db.session.commit()
    flash(f'קוד כניסה חדש למוסד: {inst.login_code}')
    return redirect(url_for('admin.institution_detail', inst_id=inst.id))
