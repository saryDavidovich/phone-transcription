"""
לשונית 'ניהול תלמידים'. תלמיד הוא בפועל שורת Customer עם institution_id
ממולא (ראה models.py) - כך שכל צנרת ההקלטה/תמלול/חיוב הקיימת עובדת עליו
בלי כפילות קוד.
"""
import os
import io
import random
import string
import threading
import logging
from datetime import datetime
from flask import Blueprint, render_template, request, jsonify, send_file, flash, redirect, url_for
from flask_login import current_user
from app import db
from models import Customer, Recording, Transaction, InstitutionUpload
from routes.institution import institution_login_required

log = logging.getLogger(__name__)

institution_students_bp = Blueprint('institution_students', __name__)


def _generate_student_number():
    for _ in range(20):
        candidate = ''.join(random.choices(string.digits, k=6))
        if not Customer.query.filter_by(student_number=candidate).first():
            return candidate
    raise RuntimeError('לא ניתן היה ליצור מספר תלמיד פנוי')


@institution_students_bp.route('/institution/students')
@institution_login_required
def students_tab():
    students = Customer.query.filter_by(institution_id=current_user.id).order_by(Customer.created_at.desc()).all()
    return render_template('institution/students.html', students=students)


@institution_students_bp.route('/institution/students/add', methods=['POST'])
@institution_login_required
def add_student():
    name = (request.form.get('name') or '').strip()
    phone = (request.form.get('phone') or '').strip() or None
    email = (request.form.get('email') or '').strip() or None
    fax = (request.form.get('fax') or '').strip() or None
    initial_balance = float(request.form.get('initial_balance') or 0)
    inst = current_user

    if phone and Customer.query.filter_by(phone=phone).first():
        flash('מספר הטלפון הזה כבר קיים במערכת')
        return redirect(url_for('institution_students.students_tab'))

    if initial_balance > 0 and (inst.balance or 0) < initial_balance:
        flash(f'אין למוסד מספיק יתרה לזיכוי הפתיחה שביקשת (יתרת המוסד: {inst.balance or 0:.2f} ₪). נא לטעון יתרה בלשונית "הגדרות חיוב", או להוסיף את התלמיד בלי זיכוי פתיחה.')
        return redirect(url_for('institution_students.students_tab'))

    student = Customer(
        name=name, phone=phone, email=email, fax=fax,
        balance=initial_balance,
        institution_id=current_user.id,
        student_number=_generate_student_number(),
        student_display_name=name,
    )
    db.session.add(student)
    db.session.commit()

    if initial_balance:
        if initial_balance > 0:
            inst.balance = (inst.balance or 0) - initial_balance
        db.session.add(Transaction(customer_id=student.id, amount=initial_balance, type='credit', description='זיכוי פתיחה ע"י המוסד'))
        db.session.commit()

    flash(f'התלמיד נוסף בהצלחה - מספר תלמיד: {student.student_number}')
    return redirect(url_for('institution_students.students_tab'))


@institution_students_bp.route('/institution/students/<int:student_id>/credit', methods=['POST'])
@institution_login_required
def credit_student(student_id):
    student = Customer.query.filter_by(id=student_id, institution_id=current_user.id).first_or_404()
    inst = current_user
    try:
        amount = float(request.form.get('amount'))
    except (TypeError, ValueError):
        flash('סכום לא תקין')
        return redirect(url_for('institution_students.students_tab'))

    # זיכוי תלמיד (סכום חיובי) יורד בפועל מיתרת המוסד עצמה - אחרת נוצר "כסף
    # מהאוויר". חיוב תלמיד (סכום שלילי) לא נוגע ביתרת המוסד.
    if amount > 0:
        if (inst.balance or 0) < amount:
            flash(f'אין למוסד מספיק יתרה לזיכוי הזה (יתרת המוסד: {inst.balance or 0:.2f} ₪). נא לטעון יתרה בלשונית "הגדרות חיוב".')
            return redirect(url_for('institution_students.students_tab'))
        inst.balance = (inst.balance or 0) - amount

    student.balance = (student.balance or 0) + amount
    db.session.add(Transaction(
        customer_id=student.id, amount=amount,
        type='credit' if amount >= 0 else 'debit',
        description='זיכוי ע"י המוסד' if amount >= 0 else 'חיוב ע"י המוסד',
    ))
    db.session.commit()
    flash('היתרה עודכנה')
    return redirect(url_for('institution_students.students_tab'))


@institution_students_bp.route('/institution/students/<int:student_id>/block', methods=['POST'])
@institution_login_required
def toggle_block(student_id):
    student = Customer.query.filter_by(id=student_id, institution_id=current_user.id).first_or_404()
    student.is_blocked = not student.is_blocked
    db.session.commit()
    flash('התלמיד נחסם' if student.is_blocked else 'החסימה בוטלה')
    return redirect(url_for('institution_students.students_tab'))


@institution_students_bp.route('/institution/students/<int:student_id>/remove', methods=['POST'])
@institution_login_required
def remove_student(student_id):
    """'הורדת תלמיד' - מנתק אותו מהמוסד (לא מוחק פיזית) כדי לא לשבור
    היסטוריית הקלטות/תנועות שכבר קיימת עבורו."""
    student = Customer.query.filter_by(id=student_id, institution_id=current_user.id).first_or_404()
    student.institution_id = None
    student.student_number = None
    db.session.commit()
    flash('התלמיד הוסר מהמוסד')
    return redirect(url_for('institution_students.students_tab'))


@institution_students_bp.route('/institution/students/excel-template')
@institution_login_required
def download_excel_template():
    """קובץ אקסל לדוגמה עם העמודות הנכונות - נלחץ מתוך המודל 'ייבוא מאקסל'
    בעמוד ניהול תלמידים, כדי שהמוסד ידע בדיוק לפי איזה פורמט למלא."""
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'תלמידים'
    ws.append(['name', 'phone', 'email', 'balance'])
    ws.append(['ישראל ישראלי', '0501234567', 'israel@example.com', 20])
    ws.append(['דוגמה בלי טלפון', '', '', 0])
    ws.column_dimensions['A'].width = 22
    ws.column_dimensions['B'].width = 16
    ws.column_dimensions['C'].width = 26
    ws.column_dimensions['D'].width = 10
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    from routes.download_utils import render_data_uri_download_page
    return render_data_uri_download_page(
        buf.getvalue(), 'תבנית_ייבוא_תלמידים.xlsx',
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )


@institution_students_bp.route('/institution/students/upload-excel', methods=['POST'])
@institution_login_required
def upload_excel():
    file = request.files.get('excel_file')
    if not file:
        flash('יש לבחור קובץ')
        return redirect(url_for('institution_students.students_tab'))

    import openpyxl
    wb = openpyxl.load_workbook(file, data_only=True)
    ws = wb.active
    headers = [str(c.value or '').strip().lower() for c in ws[1]]

    def col(*names):
        for n in names:
            if n in headers:
                return headers.index(n)
        return None

    idx_name = col('name', 'שם')
    idx_phone = col('phone', 'טלפון')
    idx_email = col('email', 'מייל', 'אימייל')
    idx_balance = col('balance', 'יתרה')

    inst = current_user
    remaining_inst_balance = inst.balance or 0
    added = 0
    skipped_balance = 0
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not row or all(v is None for v in row):
            continue
        name = str(row[idx_name]).strip() if idx_name is not None and row[idx_name] else ''
        phone = str(row[idx_phone]).strip() if idx_phone is not None and row[idx_phone] else None
        email = str(row[idx_email]).strip() if idx_email is not None and row[idx_email] else None
        requested_balance = float(row[idx_balance]) if idx_balance is not None and row[idx_balance] else 0.0

        if phone and Customer.query.filter_by(phone=phone).first():
            continue  # דילוג על טלפון כפול, לא עוצר את כל הייבוא

        # לא נותנים לתלמיד יתרת פתיחה שהמוסד בפועל לא מכסה - התלמיד עדיין
        # מיובא, רק בלי הזיכוי (או עם חלק ממנו אם היתרה הצטמצמה באמצע הקובץ).
        balance = 0.0
        if requested_balance > 0:
            if remaining_inst_balance >= requested_balance:
                balance = requested_balance
                remaining_inst_balance -= requested_balance
            else:
                skipped_balance += 1

        student = Customer(
            name=name, phone=phone, email=email, balance=balance,
            institution_id=current_user.id,
            student_number=_generate_student_number(),
            student_display_name=name,
        )
        db.session.add(student)
        added += 1

    inst.balance = remaining_inst_balance
    db.session.commit()
    msg = f'יובאו {added} תלמידים בהצלחה'
    if skipped_balance:
        msg += f' - שימו לב: ל-{skipped_balance} מהם לא ניתנה יתרת הפתיחה שביקשתם כי יתרת המוסד לא הספיקה'
    flash(msg)
    return redirect(url_for('institution_students.students_tab'))


@institution_students_bp.route('/institution/students/export-excel')
@institution_login_required
def export_excel():
    import openpyxl
    students = Customer.query.filter_by(institution_id=current_user.id).all()

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(['מספר תלמיד', 'שם', 'טלפון', 'מייל', 'יתרה', 'סה"כ הקלטות', 'חסום'])
    for s in students:
        recording_count = Recording.query.filter_by(customer_id=s.id).count()
        ws.append([s.student_number or '', s.name or '', s.phone or '', s.email or '',
                   s.balance or 0, recording_count, 'כן' if s.is_blocked else 'לא'])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    from routes.download_utils import render_data_uri_download_page
    return render_data_uri_download_page(
        buf.getvalue(), 'תלמידים.xlsx',
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )


@institution_students_bp.route('/institution/students/<int:student_id>')
@institution_login_required
def student_detail(student_id):
    student = Customer.query.filter_by(id=student_id, institution_id=current_user.id).first_or_404()
    recordings = Recording.query.filter_by(customer_id=student.id).order_by(Recording.created_at.desc()).all()
    transactions = Transaction.query.filter_by(customer_id=student.id).order_by(Transaction.created_at.desc()).limit(30).all()
    return render_template('institution/student_detail.html', student=student, recordings=recordings, transactions=transactions)


@institution_students_bp.route('/institution/students/<int:student_id>/recording/<int:recording_id>/download')
@institution_login_required
def download_student_recording(student_id, recording_id):
    """הורדת התמלול של הקלטה בודדת של תלמיד כקובץ Word - מחליף את כפתור
    "צפייה" הקודם (שהיה מציג alert() עם הטקסט, ולא באמת עבד)."""
    student = Customer.query.filter_by(id=student_id, institution_id=current_user.id).first_or_404()
    recording = Recording.query.filter_by(id=recording_id, customer_id=student.id).first_or_404()
    if not recording.transcript:
        return 'אין עדיין תמלול להקלטה הזו', 404

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    title = doc.add_heading(f'תמלול: {student.name or ""}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for line in recording.transcript.split('\n'):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    from routes.download_utils import render_data_uri_download_page
    hebrew_name = f"תמלול - {student.name or 'תלמיד'} - {recording.created_at.strftime('%d-%m-%Y') if recording.created_at else recording.id}.docx"
    return render_data_uri_download_page(
        output.getvalue(), hebrew_name,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


@institution_students_bp.route('/institution/students/<int:student_id>/upload', methods=['POST'])
@institution_login_required
def upload_for_student(student_id):
    """העלאת קובץ נוסף לתמלול עבור תלמיד ספציפי - זהה בעקרון ללשונית
    'יצירת תמלול' הכללית, רק שהתוצאה משויכת לתלמיד (Recording אמיתי,
    כולל חיוב יתרה) ולא רק קובץ Word חד-פעמי למוסד."""
    student = Customer.query.filter_by(id=student_id, institution_id=current_user.id).first_or_404()
    file = request.files.get('audio_file')
    if not file or not file.filename:
        return jsonify({'error': 'יש לבחור קובץ'}), 400

    from routes.institution import student_usage_limit_error
    limit_error = student_usage_limit_error(student)
    if limit_error:
        return jsonify({'error': limit_error}), 403

    import uuid
    static_dir = os.path.join(os.path.dirname(__file__), '..', 'static', 'fax_tmp')
    os.makedirs(static_dir, exist_ok=True)
    ext = os.path.splitext(file.filename)[1] or '.wav'
    filename = f"student_{uuid.uuid4().hex}{ext}"
    file.save(os.path.join(static_dir, filename))

    base_url = os.environ.get('APP_BASE_URL', os.environ.get('APP_URL', '')).rstrip('/')
    rec_url = f"{base_url}/static/fax_tmp/{filename}"

    recording = Recording(
        customer_id=student.id, call_id=f"inst_{uuid.uuid4().hex[:12]}",
        status='processing', source_filename=file.filename, rec_url=rec_url,
    )
    db.session.add(recording)
    db.session.commit()

    from app import app as flask_app
    t = threading.Thread(target=_process_student_upload, args=(flask_app, recording.id, rec_url), daemon=True)
    t.start()

    return jsonify({'recordingId': recording.id})


def _process_student_upload(flask_app, recording_id, rec_url):
    with flask_app.app_context():
        recording = Recording.query.get(recording_id)
        student = Customer.query.get(recording.customer_id)
        try:
            from services.transcribe import _gemini_from_url, _send_email, _send_fax
            transcript, duration, _ = _gemini_from_url(rec_url)
            if not transcript:
                recording.status = 'error'
                db.session.commit()
                return

            # מחיר לפי מחירון המנהל הראשי (מחירון בסיסי, יחידות של 20 דקות -
            # מעוגל כלפי מעלה, בדיוק כמו בכל שאר המערכת) - לא PRICE_PER_30MIN
            # הישן שלא היה קשור בכלל למחירון האמיתי.
            import math
            from routes.admin import get_setting
            price_per_20min = float(get_setting('price_per_20min_basic', '0.90'))
            units = math.ceil((duration or 0) / 1200) or 1
            cost = round(units * price_per_20min, 2)

            recording.transcript = transcript
            recording.duration_seconds = duration
            recording.cost = cost
            recording.status = 'completed'
            student.balance = (student.balance or 0) - cost
            db.session.add(Transaction(customer_id=student.id, amount=-cost, type='charge',
                                        description='תמלול הקלטה', recording_id=recording.id))
            db.session.commit()

            institution = student.institution
            target_email = student.email or (institution.notify_email if institution else None)
            target_fax = student.fax or (institution.notify_fax if institution else None)
            try:
                if target_email:
                    _send_email(target_email, transcript, student, rec_url, duration)
                elif target_fax:
                    _send_fax(target_fax, transcript, student, duration)
            except Exception:
                log.exception('שליחת תמלול תלמיד נכשלה (התמלול עצמו הצליח)')
        except Exception as e:
            log.exception(f'Student upload {recording_id} failed')
            recording.status = 'error'
            db.session.commit()


@institution_students_bp.route('/institution/students/recording-status/<int:recording_id>')
@institution_login_required
def recording_status(recording_id):
    recording = Recording.query.get_or_404(recording_id)
    student = Customer.query.get(recording.customer_id)
    if not student or student.institution_id != current_user.id:
        return jsonify({'error': 'אין הרשאה'}), 403
    return jsonify({'status': recording.status})
