import os, requests, logging, threading, time, math
from concurrent.futures import ThreadPoolExecutor

log = logging.getLogger(__name__)

# מקסימום 10 תמלולים במקביל — מגן על rate limits של גמיני
_executor = ThreadPoolExecutor(max_workers=10)
# תור תמלול — מקסימום 10 במקביל

# תור OCR נפרד — מקסימום 6 לקוחות במקביל
_ocr_executor = ThreadPoolExecutor(max_workers=6)



def ocr_async(func, *args, **kwargs):
    waiting = _ocr_executor._work_queue.qsize()
    log.info(f"ocr_async: נכנס לתור (ממתינים בתור: {waiting})")
    return _ocr_executor.submit(func, *args, **kwargs)

def transcribe_async(call_id, rec_url, customer_id, delivery_method, delivered_to, duration_seconds, transcription_tier='basic', language='he', output_language='he'):
    waiting = _ocr_executor._work_queue.qsize()
    log.info(f"transcribe_async: call_id={call_id} נכנס לתור (ממתינים בתור: {waiting})")
    _executor.submit(
        _process,
        call_id, rec_url, customer_id, delivery_method, delivered_to,
        duration_seconds, transcription_tier, language, output_language,
    )


def _process(call_id, rec_url, customer_id, delivery_method, delivered_to, duration_seconds, transcription_tier='basic', language='he', output_language='he'):
    from app import app, db
    from models import Recording, Customer, Transaction
    with app.app_context():
        try:
            rec = Recording.query.filter_by(call_id=call_id).first()
            if rec:
                rec.status = 'transcribing'
                rec.rec_url = rec_url
                db.session.commit()

            db.session.remove()

            tier = transcription_tier

            if tier == 'premium':
                log.info(f"Using AlefBot for customer {customer_id}")
                alefbot_tier = 'premium_quality'  # תמיד מקצועי
                # שפת פלט: 'he' = תרגם לעברית, 'original' (או כל ערך אחר) = השאר בשפת ההקלטה
                translate_heb = (output_language == 'he')
                job_id, actual_duration = _alefbot_submit(rec_url, call_id, model_tier=alefbot_tier, translate_to_hebrew=translate_heb)

                if job_id:
                    db.session.remove()
                    rec = Recording.query.filter_by(call_id=call_id).first()
                    if rec:
                        rec.alefbot_job_id = job_id
                        rec.status = 'alefbot_pending'
                        if actual_duration:
                            rec.duration_seconds = actual_duration
                        db.session.commit()
                    log.info(f"AlefBot job submitted: {job_id}, waiting for webhook + polling backup")

                    # Polling כגיבוי — אם webhook לא הגיע, השרת ישאל בעצמו
                    import threading
                    def _poll_alefbot(job_id, call_id, customer_id, delivery_method, delivered_to):
                        import requests as req_poll, os as _os
                        api_key = _os.environ.get('ALEFBOT_API_KEY')
                        base = 'https://alef-bot.top/api/v1'
                        from app import app, db
                        with app.app_context():
                            # בדוק כל 30 שניות עד 30 דקות
                            for attempt in range(60):
                                time.sleep(30)
                                try:
                                    # בדוק אם כבר טופל ע"י webhook
                                    from models import Recording
                                    rec_check = Recording.query.filter_by(call_id=call_id).first()
                                    if rec_check and rec_check.status not in ('alefbot_pending',):
                                        log.info(f"AlefBot poll: job {job_id} already handled (status={rec_check.status}), stopping poll")
                                        return

                                    status_res = req_poll.get(
                                        f'{base}/transcriptions/{job_id}',
                                        headers={'Authorization': f'Bearer {api_key}'},
                                        timeout=15
                                    )
                                    job_status = status_res.json().get('status', '')
                                    log.info(f"AlefBot poll {attempt+1}/60: job={job_id} status={job_status}")

                                    if job_status == 'completed':
                                        art = req_poll.get(
                                            f'{base}/transcriptions/{job_id}/artifact?format=txt',
                                            headers={'Authorization': f'Bearer {api_key}'},
                                            timeout=30
                                        )
                                        art.raise_for_status()
                                        transcript = art.text.strip()
                                        log.info(f"AlefBot poll: transcript fetched {len(transcript)} chars")
                                        finalize_alefbot_recording(call_id, transcript)
                                        return
                                    elif job_status in ('failed', 'cancelled'):
                                        log.error(f"AlefBot poll: job {job_id} failed with status {job_status}")
                                        from models import Recording
                                        rec_err = Recording.query.filter_by(call_id=call_id).first()
                                        if rec_err:
                                            rec_err.status = 'error'
                                            db.session.commit()
                                        return
                                except Exception as poll_err:
                                    log.warning(f"AlefBot poll error attempt {attempt+1}: {poll_err}")

                            log.error(f"AlefBot poll: job {job_id} timed out after 30 minutes")

                    t = threading.Thread(
                        target=_poll_alefbot,
                        args=(job_id, call_id, customer_id, delivery_method, delivered_to),
                        daemon=True
                    )
                    t.start()
                    return
                else:
                    db.session.remove()
                    rec = Recording.query.filter_by(call_id=call_id).first()
                    if rec:
                        rec.status = 'error'
                        db.session.commit()
                    return

            else:
                # gemini או כל ברירת מחדל
                log.info(f"Using Gemini for customer {customer_id}")

                # אומדן אורך לפי גודל קובץ אם לא ידוע
                if not duration_seconds or duration_seconds == 0:
                    try:
                        head_r = requests.head(rec_url, timeout=10, allow_redirects=True)
                        content_length = int(head_r.headers.get('Content-Length', 0))
                        if content_length > 0:
                            # WAV 16kHz mono 16bit: ~32KB לשנייה, אודיו דחוס: ~1KB לשנייה
                            url_lower = rec_url.lower()
                            if 'wav' in url_lower or content_length > 10_000_000:
                                duration_seconds = content_length // 32000
                            else:
                                duration_seconds = content_length // 1000
                            log.info(f"Estimated duration from Content-Length {content_length}: {duration_seconds}s")
                    except Exception as e:
                        log.warning(f"Could not estimate duration: {e}")

                # בדיקת יתרה לפי אומדן אורך לפני תמלול
                if duration_seconds and duration_seconds > 0:
                    price_per_20min_pre = float(_get_setting('price_per_20min_basic', '0.90'))
                    units_pre = math.ceil(duration_seconds / 1200)
                    cost_pre = round(units_pre * price_per_20min_pre, 2)
                    customer_pre = Customer.query.get(customer_id)
                    if customer_pre and customer_pre.balance < cost_pre:
                        log.info(f"Insufficient balance for {call_id}: need {cost_pre}, have {customer_pre.balance}")
                        db.session.remove()
                        rec = Recording.query.filter_by(call_id=call_id).first()
                        if rec:
                            _save_pending_payment(
                                rec=rec,
                                customer=customer_pre,
                                duration_seconds=duration_seconds,
                                price_per_20min=price_per_20min_pre,
                                delivery_method=delivery_method,
                                delivered_to=delivered_to,
                                transcription_tier=transcription_tier,
                                language=language,
                                output_language=output_language,
                            )
                        return

                transcript_raw, actual_duration, is_video = _gemini_from_url(rec_url, language, output_language)
                transcript_fixed = transcript_raw
                price_key = 'price_per_20min_video' if is_video else 'price_per_20min_basic'
                description_tier = 'וידאו' if is_video else 'רגיל'

            if actual_duration and actual_duration > 0:
                duration_seconds = actual_duration
                log.info(f"Actual duration from file: {duration_seconds}s")

            # סגור ופתח מחדש את ה-connection — חיוני לתמלולים ארוכים שה-SSL נסגר
            for _attempt in range(5):
                try:
                    db.session.remove()
                    db.engine.dispose()
                    break
                except Exception as _e:
                    log.warning(f"DB dispose attempt {_attempt+1} failed: {_e}")
                    time.sleep(2)

            rec = None
            for _attempt in range(5):
                try:
                    rec = Recording.query.filter_by(call_id=call_id).first()
                    break
                except Exception as _e:
                    log.warning(f"DB query attempt {_attempt+1} failed: {_e}")
                    db.session.remove()
                    db.engine.dispose()
                    time.sleep(3)

            if not transcript_raw:
                if rec:
                    rec.status = 'error'
                    db.session.commit()
                return

            price_per_20min = float(_get_setting(price_key, '0.90'))
            units = math.ceil(duration_seconds / 1200)
            cost = round(units * price_per_20min, 2)

            # בדיקה סופית שהיתרה מספיקה (לאחר שיודעים את האורך המדויק)
            customer = Customer.query.get(customer_id)
            if customer and customer.balance < cost:
                log.info(f"Insufficient balance after transcription for {call_id}: need {cost}, have {customer.balance}")
                rec = Recording.query.filter_by(call_id=call_id).first()
                if rec:
                    _save_pending_payment(
                        rec=rec,
                        customer=customer,
                        duration_seconds=duration_seconds,
                        price_per_20min=price_per_20min,
                        delivery_method=delivery_method,
                        delivered_to=delivered_to,
                        transcription_tier=transcription_tier,
                        language=language,
                        output_language=output_language,
                        transcript=transcript_fixed,  # שמור טרנסקריפט — לא לתמלל שוב
                    )
                return

            if rec:
                rec.transcript = transcript_fixed
                rec.summary = ''
                rec.status = 'transcribed'
                rec.cost = cost
                rec.duration_seconds = duration_seconds
                db.session.commit()

            if customer:
                customer.balance -= cost
                txn = Transaction(
                    customer_id=customer_id,
                    amount=-cost,
                    type='debit',
                    description=f'תמלול {duration_seconds//60} דקות ({description_tier})',
                    recording_id=rec.id if rec else None
                )
                db.session.add(txn)
                db.session.commit()

            if delivery_method == 'email':
                source_filename = rec.source_filename if rec else None
                _send_email(delivered_to, transcript_fixed, customer, rec_url, duration_seconds, source_filename=source_filename)
            elif delivery_method == 'fax':
                _send_fax(delivered_to, transcript_fixed, customer, duration_seconds, call_id)

            if rec:
                rec.status = 'delivered'
                db.session.commit()

        except Exception as e:
            log.error(f"Error processing {call_id}: {e}")


def _get_pending_recordings(customer_id):
    """מחזיר הקלטות ממתינות לתשלום שטרם פגה תקופתן"""
    from models import Recording
    from datetime import datetime
    return Recording.query.filter_by(
        customer_id=customer_id,
        status='pending_payment'
    ).filter(
        Recording.expires_at > datetime.utcnow()
    ).all()


def process_pending_recordings(customer_id):
    """נקרא אחרי טעינת ארנק - מתמלל הקלטות ממתינות אם יש יתרה מספקת"""
    from app import app, db
    from models import Recording, Customer
    from datetime import datetime

    with app.app_context():
        customer = Customer.query.get(customer_id)
        if not customer:
            return

        pending = _get_pending_recordings(customer_id)
        if not pending:
            log.info(f"process_pending_recordings: no pending for customer {customer_id}")
            return

        for rec in pending:
            price_per_20min = float(_get_setting('price_per_20min_basic', '0.90'))
            units = math.ceil((rec.duration_seconds or 0) / 1200)
            cost = round(units * price_per_20min, 2) if units > 0 else price_per_20min

            db.session.refresh(customer)
            if customer.balance < cost:
                log.info(f"process_pending_recordings: not enough balance for rec {rec.id} (need {cost}, have {customer.balance})")
                continue

            log.info(f"process_pending_recordings: processing rec {rec.id} for customer {customer_id}")

            # אם יש כבר טרנסקריפט שמור — לא לתמלל שוב, רק לחייב ולשלוח
            if rec.transcript:
                log.info(f"process_pending_recordings: transcript already exists, charging and sending")
                customer.balance -= cost
                from models import Transaction
                txn = Transaction(
                    customer_id=customer_id,
                    amount=-cost,
                    type='debit',
                    description=f'תמלול {(rec.duration_seconds or 0)//60} דקות',
                    recording_id=rec.id
                )
                db.session.add(txn)
                rec.status = 'transcribed'
                rec.cost = cost
                db.session.commit()

                # שלח למייל/פקס
                from services.transcribe import _send_email, _send_fax
                if rec.delivery_method == 'email' and rec.delivered_to:
                    _send_email(rec.delivered_to, rec.transcript, customer, rec.rec_url, rec.duration_seconds or 0)
                elif rec.delivery_method == 'fax' and rec.delivered_to:
                    _send_fax(rec.delivered_to, rec.transcript, customer, rec.duration_seconds or 0, rec.call_id)
                continue

            # אין טרנסקריפט — תמלל מחדש
            rec.status = 'queued'
            db.session.commit()

            transcribe_async(
                call_id=rec.call_id,
                rec_url=rec.rec_url,
                customer_id=customer_id,
                delivery_method=rec.delivery_method,
                delivered_to=rec.delivered_to,
                duration_seconds=rec.duration_seconds or 0,
                transcription_tier=rec.transcription_tier or 'gemini',
                language=rec.language or 'he',
                output_language=rec.output_language or 'he',
            )


def _save_pending_payment(rec, customer, duration_seconds, price_per_20min,
                           delivery_method, delivered_to, transcription_tier,
                           language, output_language, transcript=None):
    """שומר הקלטה במצב pending_payment ושולח מייל ללקוח"""
    from datetime import datetime, timedelta

    units = math.ceil(duration_seconds / 1200) if duration_seconds > 0 else 1
    cost = round(units * price_per_20min, 2)

    rec.status = 'pending_payment'
    rec.duration_seconds = duration_seconds
    rec.delivery_method = delivery_method
    rec.delivered_to = delivered_to
    rec.transcription_tier = transcription_tier
    rec.language = language
    rec.output_language = output_language
    rec.expires_at = datetime.utcnow() + timedelta(hours=72)
    # שמור טרנסקריפט אם כבר תומלל — כדי לא לתמלל שוב
    if transcript:
        rec.transcript = transcript

    from app import db
    db.session.commit()

    log.info(f"Recording {rec.call_id} saved as pending_payment, expires {rec.expires_at}, cost={cost}")

    if customer and customer.email:
        try:
            _send_insufficient_balance_email(customer.email, duration_seconds, cost, customer.balance)
        except Exception as e:
            log.error(f"Failed to send insufficient balance email: {e}")


def _send_insufficient_balance_email(to_email, duration_seconds, cost, balance):
    """שולח מייל ללקוח שהיתרה אינה מספיקה"""
    import sendgrid
    from sendgrid.helpers.mail import Mail

    minutes = duration_seconds // 60
    balance_str = f"\u20aa{balance:.2f}"
    cost_str = f"\u20aa{cost:.2f}"

    html = f"""<div dir="rtl" style="font-family:Arial,sans-serif;max-width:640px;margin:auto;color:#111827">
<h2 style="color:#dc2626">לא ניתן היה להשלים את התמלול</h2>
<p>שלום,</p>
<p>התקבל קובץ אודיו/וידאו לתמלול באורך <b>{minutes} דקות</b>.</p>
<div style="background:#fef2f2;border-right:4px solid #ef4444;padding:14px;margin:14px 0;border-radius:8px">
<p style="margin:0;font-weight:700;color:#991b1b">היתרה בארנק אינה מספיקה לתמלול זה</p>
<p style="margin:8px 0 0;color:#111827">
יתרה נוכחית: <b>{balance_str}</b><br>
עלות התמלול: <b>{cost_str}</b>
</p>
</div>
<p>הקובץ <b>נשמר במערכת ל-72 שעות</b>.<br>
אם תטעין את הארנק תוך 72 שעות, התמלול יבוצע אוטומטית ויישלח אליך.</p>
<p style="text-align:center;margin:24px 0;font-size:16px;font-weight:700;color:#1d4ed8">
לטעינת ארנק יש להתקשר למערכת ולבחור בתפריט הראשי בטעינת ארנק
</p>
<p style="color:#6b7280;font-size:13px">מערכת תמלולפון 03-3131795</p>
</div>"""

    sg = sendgrid.SendGridAPIClient(api_key=os.environ.get('SENDGRID_API_KEY'))
    message = Mail(
        from_email=os.environ.get('SENDGRID_FROM_EMAIL', os.environ.get('GMAIL_USER', '')),
        to_emails=to_email,
        subject='תמלולפון - יתרה אינה מספיקה, הקובץ נשמר',
        html_content=html,
    )
    sg.send(message)
    log.info(f"Insufficient balance email sent to {to_email}")


def _alefbot_submit(rec_url, call_id, model_tier='premium_quality', translate_to_hebrew=False):
    """שולח ל-AlefBot ומחזיר job_id מיד"""
    try:
        import wave, io

        api_key = os.environ.get('ALEFBOT_API_KEY')
        base_url = 'https://alef-bot.top/api/v1'
        webhook_url = os.environ.get('APP_BASE_URL', '').rstrip('/') + '/api/alefbot-webhook'

        r = requests.get(rec_url, timeout=300)
        r.raise_for_status()
        file_bytes = r.content
        log.info(f"Downloaded {len(file_bytes)} bytes for AlefBot")

        actual_duration = 0
        try:
            with wave.open(io.BytesIO(file_bytes)) as wav_in:
                actual_duration = wav_in.getnframes() // wav_in.getframerate()
            log.info(f"Duration: {actual_duration}s")
        except Exception as e:
            log.warning(f"Could not read WAV metadata: {e}")

        # שלב 1 — צור upload slot
        upload_res = requests.post(
            f'{base_url}/uploads',
            headers={'Authorization': f'Bearer {api_key}', 'Content-Type': 'application/json'},
            json={'filename': f'{call_id}.wav', 'content_type': 'audio/wav', 'size_bytes': len(file_bytes)},
            timeout=30
        )
        upload_res.raise_for_status()
        upload_id = upload_res.json().get('upload_id')
        log.info(f"AlefBot upload slot created: {upload_id}")

        # שלב 2 — העלה את הקובץ
        put_res = requests.put(
            f'{base_url}/uploads/{upload_id}/binary',
            headers={'Authorization': f'Bearer {api_key}'},
            data=file_bytes,
            timeout=300
        )
        put_res.raise_for_status()
        log.info(f"AlefBot file uploaded")

        # שלב 3 — צור תמלול עם webhook
        transcribe_res = requests.post(
            f'{base_url}/transcriptions',
            headers={'Authorization': f'Bearer {api_key}', 'Content-Type': 'application/json'},
            json={
                'upload_id': upload_id,
                'output_format': 'plain_text',
                'webhook_url': webhook_url,
                'model_tier': model_tier,
                'translate_to_hebrew': translate_to_hebrew,
            },
            timeout=30
        )
        if not transcribe_res.ok:
            log.error(f"AlefBot transcription error {transcribe_res.status_code}: {transcribe_res.text}")
        transcribe_res.raise_for_status()
        job_id = transcribe_res.json().get('job_id') or transcribe_res.json().get('id')
        log.info(f"AlefBot job created: {job_id} for call {call_id}")
        return job_id, actual_duration

    except Exception as e:
        log.error(f"AlefBot submit error: {e}")
        return None, 0


def finalize_alefbot_recording(call_id, transcript_text):
    """נקרא מה-webhook כשאלף בוט מסיים"""
    from app import app, db
    from models import Recording, Customer, Transaction

    with app.app_context():
        try:
            db.session.remove()
            rec = Recording.query.filter_by(call_id=call_id).first()
            if not rec:
                log.error(f"AlefBot webhook: recording not found for call {call_id}")
                return

            duration_seconds = rec.duration_seconds or 0
            price_per_20min = float(_get_setting('price_per_20min_premium', '1.90'))
            units = math.ceil(duration_seconds / 1200) if duration_seconds > 0 else 1
            cost = round(units * price_per_20min, 2)

            rec.transcript = transcript_text
            rec.summary = ''
            rec.status = 'transcribed'
            rec.cost = cost
            rec.alefbot_job_id = None
            db.session.commit()

            customer = Customer.query.get(rec.customer_id)
            if customer:
                customer.balance -= cost
                txn = Transaction(
                    customer_id=rec.customer_id,
                    amount=-cost,
                    type='debit',
                    description=f'תמלול {duration_seconds//60} דקות (מקצועי)',
                    recording_id=rec.id
                )
                db.session.add(txn)
                db.session.commit()

            rec_url = rec.rec_url or ''
            if rec.delivery_method == 'email':
                _send_email(rec.delivered_to, transcript_text, customer, rec_url, duration_seconds, source_filename=rec.source_filename)
            elif rec.delivery_method == 'fax':
                _send_fax(rec.delivered_to, transcript_text, customer, duration_seconds, call_id)

            rec.status = 'delivered'
            db.session.commit()
            log.info(f"AlefBot recording {call_id} finalized and delivered")

        except Exception as e:
            log.error(f"AlefBot finalize error for {call_id}: {e}")


def _split_wav_chunks(audio_content, mime_type, chunk_seconds=300):
    """
    מפצל audio_content לחלקים של chunk_seconds שניות.
    אם הקובץ הוא WAV - מפצל frame by frame.
    אם הקובץ הוא וידאו (לא WAV) - מחזיר רשימה עם הקובץ המקורי (אין פיצול).
    מחזיר רשימת bytes, כל אחד קובץ WAV תקין.
    """
    import wave, io

    if mime_type.startswith('video/'):
        # וידאו - לא ניתן לפצל בקלות, נחזיר כמו שזה
        log.info("Video file - no chunking, returning as-is")
        return [audio_content]

    try:
        with wave.open(io.BytesIO(audio_content)) as wav_in:
            nchannels = wav_in.getnchannels()
            sampwidth = wav_in.getsampwidth()
            framerate = wav_in.getframerate()
            total_frames = wav_in.getnframes()
            all_frames = wav_in.readframes(total_frames)

        frames_per_chunk = framerate * chunk_seconds
        total_duration = total_frames // framerate
        log.info(f"Splitting {total_duration}s audio into {chunk_seconds}s chunks")

        chunks = []
        offset = 0  # offset בבייטים
        bytes_per_frame = nchannels * sampwidth

        while offset < len(all_frames):
            chunk_frame_bytes = frames_per_chunk * bytes_per_frame
            chunk_data = all_frames[offset:offset + chunk_frame_bytes]
            offset += chunk_frame_bytes

            buf = io.BytesIO()
            with wave.open(buf, 'wb') as wav_out:
                wav_out.setnchannels(nchannels)
                wav_out.setsampwidth(sampwidth)
                wav_out.setframerate(framerate)
                wav_out.writeframes(chunk_data)
            chunks.append(buf.getvalue())

        return chunks

    except Exception as e:
        log.warning(f"Could not split WAV into chunks: {e}, returning as single chunk")
        return [audio_content]


def _gemini_from_url(url, language='he', output_language='he'):
    log.info(f"Gemini: language={language}, output_language={output_language}")
    try:
        import wave, audioop, io
        from google import genai
        from google.genai import types as gtypes

        api_key = os.environ.get('GOOGLE_API_KEY')
        client = genai.Client(api_key=api_key)

        r = requests.get(url, timeout=300)
        r.raise_for_status()
        log.info(f"Downloaded {len(r.content)} bytes for Gemini")

        actual_duration = 0
        audio_content = r.content
        file_size = len(audio_content)

        # זיהוי סוג הקובץ — לפי magic bytes ואחר כך לפי URL
        def _is_mp4(data):
            # MP4/MOV/M4A מתחילים ב-ftyp box לאחר 4 בייטים של גודל
            return len(data) > 8 and data[4:8] in (b'ftyp', b'moov', b'mdat', b'free', b'skip')

        url_lower = url.lower().split('?')[0]
        is_video_url = any(url_lower.endswith(ext) for ext in ('.mp4', '.mov', '.avi', '.mkv', '.3gp', '.m4v'))

        if _is_mp4(audio_content) or is_video_url:
            mime_type = 'video/mp4'
            ext = '.mp4'
        else:
            mime_type = 'audio/wav'
            ext = '.wav'

        # המרת וידאו/אודיו לא-WAV ל-WAV לצורך פיצול ותמלול
        if mime_type == 'video/mp4' or (mime_type == 'audio/wav' and not audio_content[:4] == b'RIFF'):
            try:
                from pydub import AudioSegment
                import tempfile, os as _os
                log.info(f"Converting {mime_type} to WAV for chunking...")
                with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp_in:
                    tmp_in.write(audio_content)
                    tmp_in_path = tmp_in.name
                try:
                    seg = AudioSegment.from_file(tmp_in_path)
                    seg = seg.set_frame_rate(16000).set_channels(1).set_sample_width(2)
                    buf = io.BytesIO()
                    seg.export(buf, format='wav')
                    audio_content = buf.getvalue()
                    actual_duration = len(seg) // 1000
                    mime_type = 'audio/wav'
                    ext = '.wav'
                    log.info(f"Converted to WAV: {len(audio_content)} bytes, {actual_duration}s")
                finally:
                    _os.unlink(tmp_in_path)
            except Exception as e:
                log.warning(f"Could not convert to WAV: {e}, sending as-is to Gemini")

        if mime_type == 'audio/wav':
            try:
                with wave.open(io.BytesIO(audio_content)) as wav_in:
                    frames = wav_in.readframes(wav_in.getnframes())
                    sampwidth = wav_in.getsampwidth()
                    nchannels = wav_in.getnchannels()
                    framerate = wav_in.getframerate()
                    actual_duration = wav_in.getnframes() // framerate

                log.info(f"Duration: {actual_duration}s, framerate: {framerate}Hz")

                if framerate != 16000:
                    frames, _ = audioop.ratecv(frames, sampwidth, nchannels, framerate, 16000, None)
                    framerate = 16000
                    log.info("Upsampled to 16000Hz for Gemini")

                    output_buffer = io.BytesIO()
                    with wave.open(output_buffer, 'wb') as wav_out:
                        wav_out.setnchannels(nchannels)
                        wav_out.setsampwidth(sampwidth)
                        wav_out.setframerate(16000)
                        wav_out.writeframes(frames)
                    audio_content = output_buffer.getvalue()
                    file_size = len(audio_content)

            except Exception as e:
                log.warning(f"Could not process WAV: {e}, using original")

        input_lang_map = {'he': 'עברית', 'yi': 'יידיש', 'en': 'English'}
        input_lang_name = input_lang_map.get(language, 'עברית')

        if output_language == 'he':
            output_instruction = 'כתוב את התמלול בעברית בלבד. אל תשתמש באותיות לטיניות.'
        elif output_language == 'yi':
            output_instruction = 'שרייב די טראנסקריפציע אויף יידיש בלעבד.'
        else:
            output_instruction = 'Write the transcription in English only.'

        if language == 'yi' and output_language == 'yi':
            yiddish_instruction = """
הדובר מדבר יידיש אשכנזית חסידית. שים לב:
- ישנם ביטויים, פסוקים וציטוטים בעברית/ארמית בהגייה אשכנזית — השאר אותם כפי שנאמרו בעברית, אל תתרגם ליידיש.
- מילים עבריות כמו "תורה", "שבת", "גמרא", "רבי" — כתוב בעברית.
- רק המשפטים שנאמרו ביידיש — כתוב ביידיש."""
        elif language == 'yi' and output_language == 'he':
            yiddish_instruction = """
הדובר מדבר יידיש אשכנזית חסידית. תרגם הכל לעברית תקנית.
ביטויים ופסוקים בעברית/ארמית — כתוב בעברית כפי שהם.
אל תשאיר אף מילה ביידיש — תרגם הכל לעברית."""
        else:
            yiddish_instruction = ''

        prompt = f"""תמלל את קובץ השמע הזה במדויק.
שפת הדובר: {input_lang_name}.
{output_instruction}
{yiddish_instruction}
חשוב ביותר — תמלול מדויק ומלא:
- תמלל כל מילה ומילה ללא יוצא מן הכלל.
- אל תדלג על אף מילה, אפילו אם הקול לא ברור — כתוב את מה שנשמע גם אם אינך בטוח.
- אל תסכם, אל תקצר, אל תדלג על חלקים.
- שמור על מינוח תורני נכון, ארמית, ראשי תיבות וגרסאות.
- החזר רק את הטקסט המתומלל ללא הערות נוספות."""

        # פיצול לחלקי 5 דקות ותמלול כל חלק בנפרד
        CHUNK_SECONDS = 300  # 5 דקות
        chunks = _split_wav_chunks(audio_content, mime_type, CHUNK_SECONDS)
        log.info(f"Split into {len(chunks)} chunks of up to {CHUNK_SECONDS}s each")

        transcript_parts = []
        for i, chunk_bytes in enumerate(chunks):
            chunk_num = i + 1
            log.info(f"Transcribing chunk {chunk_num}/{len(chunks)} ({len(chunk_bytes)} bytes)")
            chunk_transcript = None
            for attempt in range(5):
                try:
                    response = client.models.generate_content(
                        model='gemini-3.5-flash',
                        contents=[
                            prompt,
                            gtypes.Part.from_bytes(data=chunk_bytes, mime_type='audio/wav'),
                        ],
                    )
                    chunk_transcript = response.text.strip()
                    log.info(f"Chunk {chunk_num} done, {len(chunk_transcript)} chars")
                    break
                except Exception as ge:
                    log.warning(f"Chunk {chunk_num} attempt {attempt+1} failed: {ge}")
                    if attempt < 4:
                        time.sleep(15)
                    else:
                        log.error(f"Chunk {chunk_num} failed after 5 attempts, skipping")
            if chunk_transcript:
                transcript_parts.append(chunk_transcript)

        transcript = '\n\n'.join(transcript_parts) if transcript_parts else None
        if transcript:
            log.info(f"All chunks merged, total {len(transcript)} chars")

        if language == 'yi' and output_language == 'he':
            log.info("Translating Yiddish to Hebrew...")
            for attempt in range(5):
                try:
                    translate_response = client.models.generate_content(
                        model='gemini-3.5-flash',
                        contents=[f"""תרגם את הטקסט הבא מיידיש לעברית תקנית.
ביטויים ופסוקים בעברית/ארמית — השאר כפי שהם.
החזר רק את הטקסט המתורגם ללא הערות.

טקסט לתרגום:
{transcript}"""],
                    )
                    transcript = translate_response.text.strip()
                    log.info(f"Translation completed, {len(transcript)} chars")
                    break
                except Exception as te:
                    log.warning(f"Translation attempt {attempt+1} failed: {te}")
                    if attempt < 4:
                        time.sleep(10)
                    else:
                        log.error("Translation failed after 5 attempts, using original")

        return transcript, actual_duration, mime_type.startswith('video/')

    except Exception as e:
        log.error(f"Gemini error: {e}")
        return None, 0, False


def _gemini_pro_solo(url, language='he', output_language='he'):
    """
    גרסה נסיונית - תמלול בפעימה אחת בלבד באמצעות Gemini 3.1 Pro (gemini-3.1-pro-preview),
    עם פרומפט מורחב שמדגיש זיהוי הגייה אשכנזית-חסידית של עברית/ארמית.

    למשל: דובר אומר "בוריך אתוה" (הגייה אשכנזית) -> יש לתמלל "ברוך אתה" (כתיב עברי תקני),
    ולא להעתיק את ההגייה כפי שנשמעת.

    נגיש רק דרך ממשק הניהול (בדיקת תמלול), לא דרך ה-IVR או המייל.

    מחזיר: (transcript, actual_duration)
    """
    log.info(f"Gemini Pro solo: language={language}, output_language={output_language}")
    try:
        from google import genai
        from google.genai import types as gtypes

        api_key = os.environ.get('GOOGLE_API_KEY')
        client = genai.Client(api_key=api_key)

        audio_content, actual_duration = _download_and_prepare_audio(url)

        input_lang_map = {'he': 'עברית', 'yi': 'יידיש', 'en': 'English'}
        input_lang_name = input_lang_map.get(language, 'עברית')

        if output_language == 'he':
            output_instruction = 'כתוב את התמלול בעברית בלבד. אל תשתמש באותיות לטיניות.'
        elif output_language == 'yi':
            output_instruction = 'שרייב די טראנסקריפציע אויף יידיש בלעבד.'
        else:
            output_instruction = 'Write the transcription in English only.'

        if language == 'yi' and output_language == 'yi':
            yiddish_instruction = """
הדובר מדבר יידיש אשכנזית חסידית. שים לב:
- ישנם ביטויים, פסוקים וציטוטים בעברית/ארמית בהגייה אשכנזית — השאר אותם כפי שנאמרו בעברית, אל תתרגם ליידיש.
- מילים עבריות כמו "תורה", "שבת", "גמרא", "רבי" — כתוב בעברית.
- רק המשפטים שנאמרו ביידיש — כתוב ביידיש."""
        elif language == 'yi' and output_language == 'he':
            yiddish_instruction = """
הדובר מדבר יידיש אשכנזית חסידית. תרגם הכל לעברית תקנית.
ביטויים ופסוקים בעברית/ארמית — כתוב בעברית כפי שהם.
אל תשאיר אף מילה ביידיש — תרגם הכל לעברית."""
        else:
            yiddish_instruction = ''

        pronunciation_instruction = """
הוראה קריטית - זיהוי הגייה אשכנזית-חסידית של עברית וארמית:
הדובר/השומעים עשויים לצטט פסוקים, תפילות, משניות, גמרא וביטויים חז"ליים בהגייה אשכנזית-חסידית מסורתית,
שנשמעת אחרת מהעברית הישראלית המודרנית. עליך לזהות זאת ולתמלל בכתיב העברי/ארמי התקני והמדויק
של אותו מקור (לא לפי איך שזה נשמע, אלא לפי איך שזה כתוב במקור).

דוגמאות להמרת הגייה אשכנזית -> כתיב תקני (חובה לזהות תבניות כאלה ולתקן):
- "בוריך אתוה" / "בורוך אתאה" -> "ברוך אתה"
- "אדוינוי" / "אדונוי" -> "אדני" (כפי שמופיע בנוסח, ולא לכתוב "ה' " אם המקור כותב את השם המלא)
- "אלוקיינו" / "אלוקיינו מלך העוילום" -> "אלוקינו מלך העולם" (לפי הניקוד/כתיב המקובל בסידור)
- "כשם שעוסו ניסים" -> "כשם שעשו נסים"
- "תוירה" / "תויירה" -> "תורה"
- תנועות "וי"/"וא" שמייצגות "ה"/"ו" קמוצה/חולם בהגייה אשכנזית (כגון "שולוים" -> "שלום", "כוחיל" -> "כחל" וכו') - תמלל לפי הכתיב העברי הנכון של המילה, לא לפי התעתיק הפונטי.
- ת' רפה שנשמעת כ-"ס" (כגון "שבת" -> נשמע "שבס") - תמלל "שבת" (הכתיב הנכון), לא "שבס".

חשוב: זה חל **רק** על קטעי לשון הקודש (עברית/ארמית) שמצוטטים בתוך הדיבור (פסוקים, ברכות, מאמרי חז"ל,
שמות ומושגים תורניים) - לא על דיבור חולין רגיל. במקרה של ספק, בחר את הכתיב התקני/המקורי המוכר
של אותו מקור (כפי שהוא מופיע בתנ"ך/משנה/גמרא/סידור), ולא תעתיק פונטי של ההגייה."""

        prompt = f"""תמלל את קובץ השמע הזה במדויק.
שפת הדובר: {input_lang_name}.
{output_instruction}
{yiddish_instruction}
{pronunciation_instruction}
חשוב ביותר — תמלול מדויק ומלא:
- תמלל כל מילה ומילה ללא יוצא מן הכלל.
- אל תדלג על אף מילה, אפילו אם הקול לא ברור — כתוב את מה שנשמע גם אם אינך בטוח.
- אל תסכם, אל תקצר, אל תדלג על חלקים.
- שמור על מינוח תורני נכון, ארמית, ראשי תיבות וגרסאות.
- החזר רק את הטקסט המתומלל ללא הערות נוספות."""

        transcript = None
        for attempt in range(5):
            try:
                response = client.models.generate_content(
                    model='gemini-3.1-pro-preview',
                    contents=[
                        prompt,
                        gtypes.Part.from_bytes(data=audio_content, mime_type='audio/wav'),
                    ],
                )
                transcript = response.text.strip()
                log.info(f"Gemini Pro solo transcription completed, {len(transcript)} chars")
                break
            except Exception as ge:
                log.warning(f"Gemini Pro solo attempt {attempt+1} failed: {ge}")
                if attempt < 4:
                    time.sleep(15)
                else:
                    raise

        if language == 'yi' and output_language == 'he' and transcript:
            log.info("Translating Yiddish to Hebrew (Pro solo)...")
            for attempt in range(5):
                try:
                    translate_response = client.models.generate_content(
                        model='gemini-3.1-pro-preview',
                        contents=[f"""תרגם את הטקסט הבא מיידיש לעברית תקנית.
ביטויים ופסוקים בעברית/ארמית — השאר כפי שהם.
החזר רק את הטקסט המתורגם ללא הערות.

טקסט לתרגום:
{transcript}"""],
                    )
                    transcript = translate_response.text.strip()
                    log.info(f"Translation completed (Pro solo), {len(transcript)} chars")
                    break
                except Exception as te:
                    log.warning(f"Translation attempt {attempt+1} failed (Pro solo): {te}")
                    if attempt < 4:
                        time.sleep(10)
                    else:
                        log.error("Translation failed after 5 attempts (Pro solo), using original")

        return transcript, actual_duration

    except Exception as e:
        log.error(f"Gemini Pro solo error: {e}")
        return None, 0


def _download_and_prepare_audio(url):
    """מוריד את קובץ האודיו ומבצע upsample ל-16kHz אם צריך. מחזיר (audio_bytes, actual_duration)."""
    import wave, audioop, io

    r = requests.get(url, timeout=300)
    r.raise_for_status()
    audio_content = r.content
    actual_duration = 0

    try:
        with wave.open(io.BytesIO(r.content)) as wav_in:
            frames = wav_in.readframes(wav_in.getnframes())
            sampwidth = wav_in.getsampwidth()
            nchannels = wav_in.getnchannels()
            framerate = wav_in.getframerate()
            actual_duration = wav_in.getnframes() // framerate

        if framerate != 16000:
            frames, _ = audioop.ratecv(frames, sampwidth, nchannels, framerate, 16000, None)
            output_buffer = io.BytesIO()
            with wave.open(output_buffer, 'wb') as wav_out:
                wav_out.setnchannels(nchannels)
                wav_out.setsampwidth(sampwidth)
                wav_out.setframerate(16000)
                wav_out.writeframes(frames)
            audio_content = output_buffer.getvalue()
    except Exception as e:
        log.warning(f"could not process WAV: {e}, using original")

    return audio_content, actual_duration


def _gemini_review_pass(url, language='he', output_language='he'):
    """
    גרסה נסיונית - 'תמלול מקצועי' חדש המבוסס על Gemini בלבד (במקום אלף בוט).

    שלב 1: תמלול ראשוני רגיל (כמו ב-_gemini_from_url).
    שלב 2: Gemini מאזין לקובץ האודיו שוב, **תוך כדי קריאת התמלול הראשוני**,
            ומתבקש לאתר ולתקן טעויות - מילים שלא הובנו נכון, שמות, מינוחים
            תורניים, פיסוק - מבלי לשנות את התוכן/המשמעות.

    נגיש רק דרך ממשק הניהול (בדיקת תמלול), לא דרך ה-IVR או המייל,
    כך שניתן להשוות תוצאות לפני החלפת אלף בוט.

    מחזיר: (transcript_final, actual_duration, transcript_raw_first_pass)
    """
    transcript_raw, actual_duration, _ = _gemini_from_url(url, language, output_language)
    if not transcript_raw:
        return None, 0, None

    try:
        from google import genai
        from google.genai import types as gtypes

        api_key = os.environ.get('GOOGLE_API_KEY')
        client = genai.Client(api_key=api_key)

        audio_content, _ = _download_and_prepare_audio(url)

        output_lang_map = {'he': 'עברית', 'yi': 'יידיש', 'en': 'English'}
        output_lang_name = output_lang_map.get(output_language, 'עברית')

        review_prompt = f"""לפניך הקלטת שמע ותמלול ראשוני שלה (בשפה: {output_lang_name}).
האזן להקלטה במלואה תוך כדי קריאת התמלול הראשוני, והפק גרסה מתוקנת ומדויקת יותר.

הנחיות:
- תקן מילים שתומללו בצורה שגויה (כתיב, שמות, מינוחים תורניים, ראשי תיבות, ביטויים בארמית/עברית).
- תקן פיסוק וחלוקה למשפטים/פסקאות לפי מה שנשמע בפועל.
- אם יש מילים/קטעים שבתמלול הראשוני סומנו כלא ברורים או הושמטו, האזן להם שוב ונסה למלא אותם.
- אל תשנה את המשמעות, אל תסכם, אל תקצר, אל תוסיף תוכן שלא נאמר.
- אם התמלול הראשוני נכון ומדויק כפי שהוא - החזר אותו כמו שהוא, ללא שינוי מיותר.
- החזר רק את התמלול המתוקן הסופי, ללא הערות, הסברים, או ציון השינויים שנעשו.

תמלול ראשוני:
{transcript_raw}"""

        transcript_final = transcript_raw
        for attempt in range(5):
            try:
                response = client.models.generate_content(
                    model='gemini-3.5-flash',
                    contents=[
                        review_prompt,
                        gtypes.Part.from_bytes(data=audio_content, mime_type='audio/wav'),
                    ],
                )
                reviewed = response.text.strip()
                if reviewed:
                    transcript_final = reviewed
                log.info(f"Gemini review pass completed, {len(transcript_final)} chars")
                break
            except Exception as ge:
                log.warning(f"Gemini review pass attempt {attempt+1} failed: {ge}")
                if attempt < 4:
                    time.sleep(15)
                else:
                    log.error("Gemini review pass failed after 5 attempts, using first-pass transcript")

        return transcript_final, actual_duration, transcript_raw

    except Exception as e:
        log.error(f"Gemini review pass error: {e}")
        return transcript_raw, actual_duration, transcript_raw


def _gemini_dual_transcribe_and_merge(url, language='he', output_language='he', merge_model='gemini-3.5-flash'):
    """
    גרסה נסיונית - 'תמלול מקצועי' חדש המבוסס על Gemini בלבד (במקום אלף בוט).

    שלב 1: תמלול עצמאי ראשון - "שמיעה" מלאה של הקובץ (כמו _gemini_from_url).
    שלב 2: תמלול עצמאי שני - "שמיעה" נוספת ועצמאית של אותו קובץ, בלי לראות
            את התמלול הראשון, כך שלא נוצר "עיגון" לתמלול קודם.
    שלב 3: מעבר מיזוג - מודל (Flash או Pro, לפי merge_model) מאזין לאודיו
            בנוסף לשני התמלולים, ומחליט/ממזג לגרסה הסופית המדויקת ביותר.

    נגיש רק דרך ממשק הניהול (בדיקת תמלול), לא דרך ה-IVR או המייל.

    מחזיר: (transcript_final, actual_duration, transcript_a, transcript_b)
    """
    transcript_a, actual_duration, _ = _gemini_from_url(url, language, output_language)
    if not transcript_a:
        return None, 0, None, None

    # תמלול עצמאי שני - "שמיעה" נוספת ונפרדת, בלי לדעת על התמלול הראשון
    transcript_b, _, __ = _gemini_from_url(url, language, output_language)
    if not transcript_b:
        # אם התמלול השני נכשל - נמשיך עם הראשון בלבד (בלי מיזוג)
        return transcript_a, actual_duration, transcript_a, None

    try:
        from google import genai
        from google.genai import types as gtypes

        api_key = os.environ.get('GOOGLE_API_KEY')
        client = genai.Client(api_key=api_key)

        audio_content, _ = _download_and_prepare_audio(url)

        output_lang_map = {'he': 'עברית', 'yi': 'יידיש', 'en': 'English'}
        output_lang_name = output_lang_map.get(output_language, 'עברית')

        merge_prompt = f"""לפניך הקלטת שמע, ושני תמלולים עצמאיים שלה (בשפה: {output_lang_name}) -
שניהם הופקו בנפרד ע"י תמלול אוטומטי מאותו קובץ, ולכן הם עלולים להיות שונים זה מזה במקומות מסוימים.

האזן להקלטה במלואה, והשתמש בשני התמלולים כרפרנס, כדי להפיק גרסה אחת סופית - המדויקת ביותר.

הנחיות:
- במקומות שבהם שני התמלולים זהים - זה כנראה נכון, אין צורך לשנות.
- במקומות שבהם הם שונים זה מזה - האזן בקפידה לקטע המתאים בהקלטה, והכרע מי משניהם (או ניסוח שלישי) נכון יותר.
- תקן גם שגיאות נוספות שאתה שומע בהקלטה ולא מופיעות נכון באף אחד משני התמלולים (כתיב, שמות, מינוחים תורניים, ראשי תיבות, ביטויים בארמית/עברית, פיסוק).
- אל תשנה את המשמעות, אל תסכם, אל תקצר, אל תוסיף תוכן שלא נאמר בהקלטה.
- החזר רק את התמלול הסופי הממוזג, ללא הערות, הסברים, או ציון ההבדלים/ההכרעות.

תמלול א':
{transcript_a}

תמלול ב':
{transcript_b}"""

        transcript_final = transcript_a
        for attempt in range(5):
            try:
                response = client.models.generate_content(
                    model=merge_model,
                    contents=[
                        merge_prompt,
                        gtypes.Part.from_bytes(data=audio_content, mime_type='audio/wav'),
                    ],
                )
                merged = response.text.strip()
                if merged:
                    transcript_final = merged
                log.info(f"Gemini dual-transcribe merge ({merge_model}) completed, {len(transcript_final)} chars")
                break
            except Exception as ge:
                log.warning(f"Gemini merge ({merge_model}) attempt {attempt+1} failed: {ge}")
                if attempt < 4:
                    time.sleep(15)
                else:
                    log.error(f"Gemini merge ({merge_model}) failed after 5 attempts, using first transcript")

        return transcript_final, actual_duration, transcript_a, transcript_b

    except Exception as e:
        log.error(f"Gemini dual-transcribe merge error: {e}")
        return transcript_a, actual_duration, transcript_a, transcript_b


def _get_setting(key, default=''):
    from models import Settings
    s = Settings.query.filter_by(key=key).first()
    return s.value if s else default


def _build_word_doc(name, duration_str, transcript_fixed, transcript_raw=None, title='תמלול שיחה'):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    import io

    def set_rtl(paragraph, justify=False):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.RIGHT
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)

    def add_bottom_border(paragraph):
        from docx.oxml.ns import qn
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '999999')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
    def add_footer(doc):
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = footer_para._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
        run = footer_para.add_run('הופק באמצעות מערכת תמלולפון 03-3131795')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc = Document()
    section = doc.sections[0]
    sectPr = section._sectPr
    bidi_doc = OxmlElement('w:bidi')
    sectPr.append(bidi_doc)
    add_footer(doc)

    title_heading = doc.add_heading(title, 0)
    set_rtl(title_heading)
    p_info = doc.add_paragraph(f'לקוח: {name} | משך: {duration_str}')
    set_rtl(p_info)
    add_bottom_border(p_info)
    h1 = doc.add_heading('תמלול', level=1)
    set_rtl(h1)
    p = doc.add_paragraph()
    set_rtl(p, justify=True)
    run_body = p.add_run(transcript_fixed or '')
    run_body.font.size = Pt(13)
    run_body.font.name = 'David'
    # הגדרת פונט מפורשת ל-complex script (עברית) — חיוני לחדות ברינדור PDF/LibreOffice
    rPr = run_body._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:cs'), 'David')
    rFonts.set(qn('w:ascii'), 'David')
    rFonts.set(qn('w:hAnsi'), 'David')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _build_pdf_for_fax(name, duration_str, transcript_fixed):
    """
    בונה PDF עם תמלול בעברית עבור שליחת פקס.
    משתמש באותו מסמך Word שנבנה למייל (_build_word_doc, RTL תקני דרך w:bidi)
    וממיר אותו ל-PDF באמצעות LibreOffice headless.
    גישה זו מטמיעה את הפונטים כ-PDF וקטורי תקני (TrueType subset עם cmap מלא),
    ונמנעת מבעיות ריבועים שחורים / טשטוש שקרו עם reportlab + רסטריזציה.
    """
    import subprocess
    import tempfile
    import uuid

    try:
        docx_bytes = _build_word_doc(name, duration_str, transcript_fixed)

        import shutil
        if not shutil.which('soffice'):
            log.error("PDF build error: soffice (LibreOffice) not found on PATH - check nixpacks.toml")
            return None

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, f'{uuid.uuid4().hex}.docx')
            with open(docx_path, 'wb') as f:
                f.write(docx_bytes)

            result = subprocess.run(
                [
                    'soffice', '--headless', '--convert-to',
                    'pdf:writer_pdf_Export:{"Quality":{"type":"long","value":100},"ReduceImageResolution":{"type":"boolean","value":false},"SelectPdfVersion":{"type":"long","value":1}}',
                    '--outdir', tmpdir, docx_path
                ],
                capture_output=True, text=True, timeout=120
            )
            if result.returncode != 0:
                log.error(f"LibreOffice convert error: {result.stderr}")
                return None

            pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
            if not os.path.exists(pdf_path):
                log.error(f"LibreOffice did not produce PDF. stdout={result.stdout} stderr={result.stderr}")
                return None

            with open(pdf_path, 'rb') as f:
                return f.read()

    except Exception as e:
        log.error(f"PDF build error: {e}")
        return None


def _normalize_israeli_phone(raw):
    """מנקה ומנרמל מספר טלפון ישראלי לפורמט מקומי (05XXXXXXXX / 0XXXXXXXXX)."""
    phone = (raw or '').strip().replace('-', '').replace(' ', '')
    if phone.startswith('+972'):
        phone = '0' + phone[4:]
    elif phone.startswith('972'):
        phone = '0' + phone[3:]
    return phone


def _send_fax(to_number, transcript_fixed, customer, duration_seconds, call_id=None):
    """
    שולח את התמלול כפקס באמצעות ה-API של ימות המשיח (SendFax).
    מעלה את ה-PDF בעצמו (pdfFile=UPLOAD) ומבקש דוח מסירה ל-deliveryUrl,
    כדי שסטטוס השליחה יתעדכן ויוצג בממשק הניהול.
    """
    try:
        name = customer.name if hasattr(customer, 'name') and customer.name else customer.phone if customer else ''
        minutes = duration_seconds // 60
        seconds = duration_seconds % 60
        duration_str = f"{minutes}:{seconds:02d}"

        pdf_bytes = _build_pdf_for_fax(name, duration_str, transcript_fixed)
        if not pdf_bytes:
            log.error("Failed to build PDF for fax")
            return

        yemot_token = os.environ.get('YEMOT_TOKEN')
        if not yemot_token:
            log.error("YEMOT_TOKEN not configured - cannot send fax")
            return

        caller_id = os.environ.get('YEMOT_FAX_CALLER_ID', '')
        base_url = os.environ.get('APP_BASE_URL', '').rstrip('/')

        # מספר היעד הוא מה שהלקוח הזין במערכת הטלפונית (to_number)
        fax_number = _normalize_israeli_phone(to_number)

        files = {
            'fileToUpload': (f'transcript_{call_id or "fax"}.pdf', pdf_bytes, 'application/pdf'),
        }
        data = {
            'token': yemot_token,
            'pdfFile': 'UPLOAD',
            'phone': fax_number,
        }
        if caller_id:
            data['callerId'] = caller_id
        if base_url and call_id:
            data['deliveryUrl'] = f'{base_url}/api/fax-delivery-webhook'

        response = requests.post(
            'https://www.call2all.co.il/ym/api/SendFax',
            data=data,
            files=files,
            timeout=120,
        )
        result = response.json() if response.headers.get('content-type', '').startswith('application/json') else {}

        if result.get('responseStatus') == 'OK':
            campaign_id = result.get('CampaignId')
            log.info(f"Fax queued via Yemot to {fax_number}, CampaignId: {campaign_id}")
            if call_id and campaign_id:
                _save_fax_campaign(call_id, campaign_id)
        else:
            log.error(f"Yemot SendFax failed: {response.text}")
            if call_id:
                _update_fax_status(call_id, status='error', note=response.text[:500])

    except Exception as e:
        log.error(f"Fax error: {e}")
        if call_id:
            _update_fax_status(call_id, status='error', note=str(e)[:500])


def _save_fax_campaign(call_id, campaign_id):
    from app import app, db
    from models import Recording
    with app.app_context():
        try:
            db.session.remove()
            rec = Recording.query.filter_by(call_id=call_id).first()
            if rec:
                rec.fax_campaign_id = campaign_id
                rec.fax_status = 'sent'
                db.session.commit()
        except Exception as e:
            log.error(f"_save_fax_campaign error: {e}")


def _update_fax_status(call_id, status, note=''):
    from app import app, db
    from models import Recording
    with app.app_context():
        try:
            db.session.remove()
            rec = Recording.query.filter_by(call_id=call_id).first()
            if rec:
                rec.fax_status = status
                if note:
                    rec.fax_status_note = note
                db.session.commit()
        except Exception as e:
            log.error(f"_update_fax_status error: {e}")


def handle_fax_delivery_webhook(data):
    """
    מטפל ב-callback של deliveryUrl מימות המשיח עבור SendFax.
    מעדכן את סטטוס הפקס של ההקלטה המתאימה (לפי CampaignId) כדי שיוצג בממשק הניהול.

    שדות אפשריים מימות:
    - CampaignId
    - Delivery: Answer / NoAnswer / End
    - DIALSTATUS (אם Delivery=NoAnswer)
    - status (אם Delivery=End) - SUCCESS במקרה של מסירה מוצלחת
    """
    from app import app, db
    from models import Recording

    campaign_id = data.get('CampaignId', '')
    delivery = data.get('Delivery', '')
    end_status = data.get('status', '')
    dial_status = data.get('DIALSTATUS', '')

    if not campaign_id:
        return

    with app.app_context():
        try:
            db.session.remove()
            rec = Recording.query.filter_by(fax_campaign_id=campaign_id).first()
            if not rec:
                log.warning(f"Fax delivery webhook: no recording for CampaignId {campaign_id}")
                return

            if delivery == 'Answer':
                rec.fax_status = 'sending'
            elif delivery == 'NoAnswer':
                rec.fax_status = 'no_answer'
                rec.fax_status_note = dial_status
            elif delivery == 'End':
                if end_status == 'SUCCESS':
                    rec.fax_status = 'delivered'
                else:
                    rec.fax_status = 'failed'
                    rec.fax_status_note = end_status

            db.session.commit()
            log.info(f"Fax status updated for CampaignId {campaign_id}: {rec.fax_status}")

        except Exception as e:
            log.error(f"handle_fax_delivery_webhook error: {e}")


def _send_email(to, transcript, customer, rec_url, duration_seconds, source_filename=None):
    try:
        import sendgrid, base64
        from sendgrid.helpers.mail import Mail, Attachment, FileContent, FileName, FileType, Disposition

        name = customer.name if hasattr(customer, 'name') and customer.name else customer.phone if customer else ''
        minutes = duration_seconds // 60
        seconds = duration_seconds % 60
        duration_str = f"{minutes}:{seconds:02d}"

        # אם ההקלטה התקבלה במייל - הכותרת היא שם הקובץ שנשלח, ולא נכלל קישור להורדת ההקלטה
        title = source_filename if source_filename else 'תמלול שיחה'

        word_bytes = _build_word_doc(name, duration_str, transcript, title=title)
        word_b64 = base64.b64encode(word_bytes).decode('utf-8')

        download_block = ''
        if not source_filename:
            download_block = f'''<div style="background:#fff7ed;border-right:4px solid #f97316;padding:16px;margin:16px 0;border-radius:8px">
<a href="{rec_url}" style="color:#ea580c;font-weight:600;font-size:15px;text-decoration:none">⬇️ להורדת ההקלטה לחצו כאן</a>
</div>'''

        html = f'''<div dir="rtl" style="font-family:Arial,sans-serif;max-width:600px;margin:auto">
<h2 style="color:#1d4ed8">{title}</h2>
<p style="color:#6b7280">לקוח: <b>{name}</b> | משך: <b>{duration_str}</b></p>
<div style="background:#f0fdf4;border-right:4px solid #10b981;padding:16px;margin:16px 0;border-radius:8px">
<h3 style="margin:0 0 12px;color:#065f46">✨ תמלול</h3>
<div style="line-height:1.8;white-space:pre-wrap;text-align:justify">{transcript}</div>
</div>
{download_block}
</div>'''

        sg = sendgrid.SendGridAPIClient(api_key=os.environ.get('SENDGRID_API_KEY'))
        message = Mail(
            from_email=os.environ.get('SENDGRID_FROM_EMAIL', os.environ.get('GMAIL_USER', '')),
            to_emails=to,
            subject=f'תמלול שיחה - {title}' if source_filename else f'תמלול שיחה - {name}',
            html_content=html
        )
        message.attachment = Attachment(
            FileContent(word_b64),
            FileName(f'תמלול_{name}.docx'),
            FileType('application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
            Disposition('attachment')
        )
        sg.send(message)
        log.info(f"Email sent to {to}")
    except Exception as e:
        log.error(f"Email error: {e}")
